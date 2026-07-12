"""Extração de documentos DOCX para o modelo canônico do Compare-Docs.

Percorre o corpo do documento na ordem REAL do XML (``w:p`` e ``w:tbl``
intercalados), preservando formatação em nível de run, títulos, listas,
tabelas e imagens inline.
"""
from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path
from typing import List, Optional

import docx
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import Block, BlockKind, Cell, Document, Run

logger = logging.getLogger(__name__)

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
}


def _length_pt(value) -> Optional[float]:
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except (TypeError, ValueError, AttributeError):
        return None


def _doc_defaults(docx_doc) -> tuple:
    """Fonte padrão do documento (docDefaults / Normal)."""
    font_name = None
    font_size_pt = None
    try:
        normal = docx_doc.styles["Normal"]
        font_name = normal.font.name
        font_size_pt = _length_pt(normal.font.size)
    except Exception:
        pass
    if not font_name:
        try:
            r_pr = docx_doc.styles.element.xpath(
                "w:docDefaults/w:rPrDefault/w:rPr/w:rFonts/@w:ascii"
            )
            if r_pr:
                font_name = r_pr[0]
        except Exception:
            pass
    if not font_size_pt:
        try:
            sz = docx_doc.styles.element.xpath(
                "w:docDefaults/w:rPrDefault/w:rPr/w:sz/@w:val"
            )
            if sz:
                font_size_pt = round(int(sz[0]) / 2.0, 2)
        except Exception:
            pass
    return font_name, font_size_pt

_HEADING_NAME_RE = re.compile(
    r"^(?:heading|t[íi]tulo)\s+(\d+)\b", re.IGNORECASE
)
_LIST_NAME_RE = re.compile(
    r"^(?:list\b|lista\b|par[áa]grafo da lista)", re.IGNORECASE
)


# ---------------------------------------------------------------------------
# Helpers de estilo/formatação
# ---------------------------------------------------------------------------

def _style_font_attr(style, attr: str) -> Optional[bool]:
    """Busca um atributo de fonte (bold/italic/...) na cadeia de estilos."""
    seen = set()
    current = style
    while current is not None and getattr(current, "style_id", None) not in seen:
        seen.add(current.style_id)
        try:
            value = getattr(current.font, attr, None)
        except Exception:  # estilo sem fonte definida
            value = None
        if value is not None:
            return bool(value)
        try:
            current = current.base_style
        except Exception:
            current = None
    return None


def _effective_flag(run_value, style, attr: str) -> bool:
    """Formatação efetiva: valor do run; se None, herda do estilo; senão False."""
    if run_value is not None:
        return bool(run_value)
    inherited = _style_font_attr(style, attr)
    if inherited is not None:
        return inherited
    return False


def _style_font_name(style) -> Optional[str]:
    seen = set()
    current = style
    while current is not None and getattr(current, "style_id", None) not in seen:
        seen.add(current.style_id)
        try:
            name = current.font.name
            if name:
                return name
        except Exception:
            pass
        try:
            current = current.base_style
        except Exception:
            current = None
    return None


def _style_font_size_pt(style) -> Optional[float]:
    seen = set()
    current = style
    while current is not None and getattr(current, "style_id", None) not in seen:
        seen.add(current.style_id)
        try:
            pt = _length_pt(current.font.size)
            if pt is not None:
                return pt
        except Exception:
            pass
        try:
            current = current.base_style
        except Exception:
            current = None
    return None


def _normalize_font_face(name: Optional[str]) -> tuple:
    """Separa sufixos Bold/Italic embutidos no nome da fonte (comum em DOCX legais)."""
    if not name:
        return None, False, False
    bold = False
    italic = False
    n = name.strip()
    if n.endswith(" Bold Italic"):
        n = n[:-12].strip()
        bold = italic = True
    elif n.endswith(" Bold"):
        n = n[:-5].strip()
        bold = True
    elif n.endswith(" Italic"):
        n = n[:-7].strip()
        italic = True
    return (n or None), bold, italic


def _effective_font(run, style, default_font: Optional[str], default_size_pt: Optional[float]):
    name = None
    size_pt = None
    try:
        name = run.font.name
    except Exception:
        name = None
    try:
        size_pt = _length_pt(run.font.size)
    except Exception:
        size_pt = None
    if not name:
        name = _style_font_name(style)
    if not name:
        name = default_font
    name, face_bold, face_italic = _normalize_font_face(name)
    if size_pt is None:
        size_pt = _style_font_size_pt(style)
    if size_pt is None:
        size_pt = default_size_pt
    return name, size_pt, face_bold, face_italic


def _style_paragraph_align(style) -> Optional[str]:
    seen = set()
    current = style
    while current is not None and getattr(current, "style_id", None) not in seen:
        seen.add(current.style_id)
        try:
            pf = current.paragraph_format
            if pf.alignment is not None:
                return _ALIGN_MAP.get(pf.alignment, "left")
        except Exception:
            pass
        try:
            current = current.base_style
        except Exception:
            current = None
    return None


def _paragraph_layout(paragraph: DocxParagraph) -> dict:
    pf = paragraph.paragraph_format
    align = None
    try:
        if pf.alignment is not None:
            align = _ALIGN_MAP.get(pf.alignment, "left")
    except Exception:
        align = None
    style = None
    try:
        style = paragraph.style
    except Exception:
        style = None
    if align is None and style is not None:
        align = _style_paragraph_align(style)
    style_name = None
    try:
        style_name = (paragraph.style.name or "").strip() or None
    except Exception:
        style_name = None
    return {
        "style_name": style_name,
        "align": align,
        "indent_left_pt": _length_pt(pf.left_indent),
        "indent_right_pt": _length_pt(pf.right_indent),
        "indent_first_pt": _length_pt(pf.first_line_indent),
        "space_before_pt": _length_pt(pf.space_before),
        "space_after_pt": _length_pt(pf.space_after),
    }


def _apply_layout(block: Block, layout: dict) -> None:
    block.style_name = layout.get("style_name")
    block.align = layout.get("align")
    block.indent_left_pt = layout.get("indent_left_pt")
    block.indent_right_pt = layout.get("indent_right_pt")
    block.indent_first_pt = layout.get("indent_first_pt")
    block.space_before_pt = layout.get("space_before_pt")
    block.space_after_pt = layout.get("space_after_pt")


def _runs_from_paragraph(
    paragraph: DocxParagraph,
    default_font: Optional[str] = None,
    default_size_pt: Optional[float] = None,
) -> List[Run]:
    """Converte os runs de um parágrafo em Runs do modelo, herdando estilo."""
    style = None
    try:
        style = paragraph.style
    except Exception:
        style = None

    result: List[Run] = []
    for r in paragraph.runs:
        text = r.text
        if not text:
            continue
        bold = _effective_flag(r.bold, style, "bold")
        italic = _effective_flag(r.italic, style, "italic")
        underline = _effective_flag(r.underline, style, "underline")
        try:
            strike_val = r.font.strike
        except Exception:
            strike_val = None
        strike = _effective_flag(strike_val, style, "strike")
        font_name, font_size_pt, face_bold, face_italic = _effective_font(
            r, style, default_font, default_size_pt
        )
        if face_bold:
            bold = True
        if face_italic:
            italic = True

        run = Run(
            text=text, bold=bold, italic=italic,
            underline=underline, strike=strike,
            font_name=font_name, font_size_pt=font_size_pt,
        )
        # Funde runs adjacentes com a mesma formatação (menos ruído no diff).
        if result and result[-1].style_key() == run.style_key():
            result[-1].text += run.text
        else:
            result.append(run)
    return result


def _heading_level(paragraph: DocxParagraph) -> int:
    """Nível de título (1..9); 0 se o parágrafo não for título."""
    style = None
    try:
        style = paragraph.style
    except Exception:
        style = None

    name = ""
    if style is not None:
        try:
            name = style.name or ""
        except Exception:
            name = ""
    match = _HEADING_NAME_RE.match(name.strip())
    if match:
        level = int(match.group(1))
        if 1 <= level <= 9:
            return level

    # Outline level: no pPr do próprio parágrafo...
    vals: List[str] = []
    try:
        vals = paragraph._p.xpath("./w:pPr/w:outlineLvl/@w:val")
    except Exception:
        vals = []
    # ...ou na cadeia de estilos.
    if not vals and style is not None:
        seen = set()
        current = style
        while current is not None and getattr(current, "style_id", None) not in seen:
            seen.add(current.style_id)
            try:
                vals = current.element.xpath("./w:pPr/w:outlineLvl/@w:val")
            except Exception:
                vals = []
            if vals:
                break
            try:
                current = current.base_style
            except Exception:
                current = None
    if vals:
        try:
            lvl = int(vals[0])
        except (TypeError, ValueError):
            return 0
        if 0 <= lvl <= 8:
            return lvl + 1
    return 0


def _is_list_item(paragraph: DocxParagraph) -> bool:
    """Detecta item de lista via numPr no XML ou nome de estilo de lista."""
    try:
        if paragraph._p.xpath("./w:pPr/w:numPr/w:numId"):
            return True
    except Exception:
        pass
    try:
        name = (paragraph.style.name or "").strip()
    except Exception:
        name = ""
    return bool(_LIST_NAME_RE.match(name))


def _image_hashes(paragraph: DocxParagraph, part) -> List[str]:
    """SHA-1 dos bytes de cada imagem inline (w:drawing / a:blip) do parágrafo."""
    hashes: List[str] = []
    rids: List[str] = []
    try:
        rids.extend(paragraph._p.xpath(".//a:blip/@r:embed"))
        rids.extend(paragraph._p.xpath(".//a:blip/@r:link"))
    except Exception:
        rids = []
    for rid in rids:
        try:
            related = part.related_parts[rid]
            blob = related.blob
        except Exception:
            logger.warning("Imagem com relacionamento inacessível: %s", rid)
            continue
        if blob:
            hashes.append(hashlib.sha1(blob).hexdigest())
    return hashes


def _cell_to_model(cell) -> Cell:
    """Converte uma célula de tabela; parágrafos da célula unidos com \\n."""
    runs: List[Run] = []
    for paragraph in cell.paragraphs:
        para_runs = _runs_from_paragraph(paragraph)
        if not para_runs:
            continue
        if runs:
            runs.append(Run(text="\n"))
        runs.extend(para_runs)
    return Cell(runs=runs)


def _table_to_block(table: DocxTable) -> Optional[Block]:
    rows: List[List[Cell]] = []
    try:
        for row in table.rows:
            rows.append([_cell_to_model(c) for c in row.cells])
    except Exception as exc:
        logger.warning("Tabela ignorada por estrutura inesperada: %s", exc)
        return None
    if not rows:
        return None
    return Block(kind=BlockKind.TABLE, rows=rows, page=None)


# ---------------------------------------------------------------------------
# Função principal
# ---------------------------------------------------------------------------

HEADER_STYLE = "__header__"
FOOTER_STYLE = "__footer__"


# ---------------------------------------------------------------------------
# Numeração automática (numbering.xml) — resolve o rótulo efetivo "(a)", "6.1."
# de cada parágrafo numerado. Sem isso, renumerar cláusulas (a→b) é INVISÍVEL
# para a comparação, porque as letras não existem no texto dos runs.
# ---------------------------------------------------------------------------

_ROMANS = [
    (1000, "m"), (900, "cm"), (500, "d"), (400, "cd"), (100, "c"),
    (90, "xc"), (50, "l"), (40, "xl"), (10, "x"), (9, "ix"),
    (5, "v"), (4, "iv"), (1, "i"),
]


def _to_roman(n: int) -> str:
    out = []
    for value, sym in _ROMANS:
        while n >= value:
            out.append(sym)
            n -= value
    return "".join(out)


def _to_letter(n: int) -> str:
    # 1→a, 26→z, 27→aa …
    out = ""
    while n > 0:
        n, rem = divmod(n - 1, 26)
        out = chr(ord("a") + rem) + out
    return out


def _format_counter(fmt: str, value: int) -> Optional[str]:
    fmt = (fmt or "decimal").strip()
    if fmt == "decimal":
        return str(value)
    if fmt == "lowerLetter":
        return _to_letter(value)
    if fmt == "upperLetter":
        return _to_letter(value).upper()
    if fmt == "lowerRoman":
        return _to_roman(value)
    if fmt == "upperRoman":
        return _to_roman(value).upper()
    if fmt in ("bullet", "none"):
        return None
    return str(value)  # formatos exóticos: melhor aproximação


class _NumberingResolver:
    """Resolve rótulos de numeração na ordem do documento."""

    def __init__(self, docx_doc) -> None:
        self._levels: dict = {}    # (numId, ilvl) -> {fmt, lvlText, start}
        self._counters: dict = {}  # numId -> {ilvl: contador atual}
        self._style_numpr: dict = {}
        self._load(docx_doc)

    def _load(self, docx_doc) -> None:
        root = None
        try:
            for rel in docx_doc.part.rels.values():
                if rel.reltype.endswith("/numbering"):
                    root = rel.target_part._element
                    break
        except Exception:
            root = None
        if root is None:
            return
        abstract: dict = {}
        for anode in root.findall(qn("w:abstractNum")):
            aid = anode.get(qn("w:abstractNumId"))
            levels = {}
            for lvl in anode.findall(qn("w:lvl")):
                ilvl = int(lvl.get(qn("w:ilvl") or "0") or 0)
                fmt_el = lvl.find(qn("w:numFmt"))
                text_el = lvl.find(qn("w:lvlText"))
                start_el = lvl.find(qn("w:start"))
                levels[ilvl] = {
                    "fmt": fmt_el.get(qn("w:val")) if fmt_el is not None else "decimal",
                    "lvlText": text_el.get(qn("w:val")) if text_el is not None else "%1.",
                    "start": int(start_el.get(qn("w:val"))) if start_el is not None else 1,
                }
            abstract[aid] = levels
        for nnode in root.findall(qn("w:num")):
            num_id = nnode.get(qn("w:numId"))
            aref = nnode.find(qn("w:abstractNumId"))
            levels = dict(abstract.get(
                aref.get(qn("w:val")) if aref is not None else None, {}
            ))
            # lvlOverride: startOverride e/ou lvl completo
            for override in nnode.findall(qn("w:lvlOverride")):
                ilvl = int(override.get(qn("w:ilvl")) or 0)
                info = dict(levels.get(ilvl, {"fmt": "decimal", "lvlText": "%1.", "start": 1}))
                so = override.find(qn("w:startOverride"))
                if so is not None:
                    info["start"] = int(so.get(qn("w:val")) or 1)
                lvl_el = override.find(qn("w:lvl"))
                if lvl_el is not None:
                    fmt_el = lvl_el.find(qn("w:numFmt"))
                    text_el = lvl_el.find(qn("w:lvlText"))
                    if fmt_el is not None:
                        info["fmt"] = fmt_el.get(qn("w:val"))
                    if text_el is not None:
                        info["lvlText"] = text_el.get(qn("w:val"))
                levels[ilvl] = info
            for ilvl, info in levels.items():
                self._levels[(num_id, ilvl)] = info

    def _effective_numpr(self, paragraph: DocxParagraph):
        ppr = paragraph._p.find(qn("w:pPr"))
        numpr = ppr.find(qn("w:numPr")) if ppr is not None else None
        if numpr is None:
            # numeração herdada do estilo (cadeia basedOn, profundidade limitada)
            style = paragraph.style
            depth = 0
            while style is not None and depth < 8:
                key = getattr(style, "style_id", None)
                if key in self._style_numpr:
                    numpr = self._style_numpr[key]
                    break
                spr = style.element.find(qn("w:pPr"))
                cand = spr.find(qn("w:numPr")) if spr is not None else None
                if cand is not None:
                    self._style_numpr[key] = cand
                    numpr = cand
                    break
                style = getattr(style, "base_style", None)
                depth += 1
        if numpr is None:
            return None
        nid = numpr.find(qn("w:numId"))
        ilvl = numpr.find(qn("w:ilvl"))
        if nid is None:
            return None
        num_id = nid.get(qn("w:val"))
        if not num_id or num_id == "0":
            return None
        return num_id, int(ilvl.get(qn("w:val")) or 0) if ilvl is not None else 0

    def label_for(self, paragraph: DocxParagraph) -> Optional[str]:
        """Rótulo do parágrafo ("(a)", "6.1.") ou None se não numerado/bullet."""
        eff = self._effective_numpr(paragraph)
        if eff is None:
            return None
        num_id, ilvl = eff if isinstance(eff, tuple) else (eff, 0)
        info = self._levels.get((num_id, ilvl))
        if info is None:
            return None
        counters = self._counters.setdefault(num_id, {})
        counters[ilvl] = counters.get(ilvl, info["start"] - 1) + 1
        # níveis mais fundos reiniciam
        for deeper in [k for k in counters if k > ilvl]:
            del counters[deeper]
        if _format_counter(info["fmt"], 1) is None:
            return None  # bullet
        label = info["lvlText"] or "%1."
        for level in range(9):
            token = "%%%d" % (level + 1)
            if token not in label:
                continue
            lvl_info = self._levels.get((num_id, level), info)
            value = counters.get(level)
            if value is None:
                value = lvl_info.get("start", 1)
                counters[level] = value
            rendered = _format_counter(lvl_info.get("fmt", "decimal"), value)
            label = label.replace(token, rendered or "")
        return label.strip() or None


def _header_footer_blocks(docx_doc) -> List[Block]:
    """Blocos de cabeçalho/rodapé distintos e não-vazios do documento.

    Deduplicados por texto (o mesmo rodapé repetido em várias seções vira um
    bloco só) e ordenados de forma estável: cabeçalhos primeiro, depois
    rodapés — assim os dois lados da comparação alinham entre si.
    """
    collected: List[Block] = []
    seen: set = set()
    for style, parts in (
        (HEADER_STYLE, ("header", "first_page_header", "even_page_header")),
        (FOOTER_STYLE, ("footer", "first_page_footer", "even_page_footer")),
    ):
        for section in docx_doc.sections:
            for attr in parts:
                try:
                    hf = getattr(section, attr)
                    if hf.is_linked_to_previous:
                        continue
                    paragraphs = hf.paragraphs
                except Exception:
                    continue
                for paragraph in paragraphs:
                    runs = _runs_from_paragraph(paragraph)
                    text = "".join(r.text for r in runs).strip()
                    if not text:
                        continue
                    key = (style, re.sub(r"\s+", " ", text))
                    if key in seen:
                        continue
                    seen.add(key)
                    collected.append(
                        Block(
                            kind=BlockKind.PARAGRAPH,
                            runs=runs,
                            page=None,
                            style_name=style,
                        )
                    )
    return collected


def extract_docx(path: str) -> Document:
    """Extrai um arquivo .docx para o modelo canônico ``Document``.

    Levanta ``ValueError`` (mensagens em pt-BR) para arquivo inexistente,
    corrompido ou sem conteúdo extraível.
    """
    p = Path(path)
    if not p.is_file():
        raise ValueError("Arquivo não encontrado: '%s'" % path)

    try:
        docx_doc = docx.Document(str(p))
    except Exception as exc:
        raise ValueError(
            "Não foi possível abrir o DOCX '%s': arquivo inválido ou corrompido (%s)"
            % (path, exc)
        ) from exc

    part = docx_doc.part
    blocks: List[Block] = []
    default_font, default_size_pt = _doc_defaults(docx_doc)
    numbering = _NumberingResolver(docx_doc)

    for child in docx_doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = DocxParagraph(child, docx_doc)
            runs = _runs_from_paragraph(paragraph, default_font, default_size_pt)
            layout = _paragraph_layout(paragraph)
            image_hashes = _image_hashes(paragraph, part)

            # SEMPRE consumir o número (mesmo em parágrafo vazio) — os
            # contadores precisam avançar na ordem real do documento.
            try:
                list_label = numbering.label_for(paragraph)
            except Exception:
                logger.debug("Falha ao resolver numeração", exc_info=True)
                list_label = None

            has_text = any(r.text.strip() for r in runs)
            if has_text:
                level = _heading_level(paragraph)
                if level > 0:
                    kind = BlockKind.HEADING
                elif _is_list_item(paragraph):
                    kind = BlockKind.LIST_ITEM
                    level = 0
                else:
                    kind = BlockKind.PARAGRAPH
                    level = 0
                block = Block(
                    kind=kind, runs=runs, level=level, page=None,
                    list_label=list_label,
                )
                _apply_layout(block, layout)
                blocks.append(block)
            for image_hash in image_hashes:
                blocks.append(
                    Block(kind=BlockKind.IMAGE, image_hash=image_hash, page=None)
                )
            # Parágrafo totalmente vazio (sem texto e sem imagem): pula.
        elif child.tag == qn("w:tbl"):
            table_block = _table_to_block(DocxTable(child, docx_doc))
            if table_block is not None:
                blocks.append(table_block)

    if not blocks:
        raise ValueError(
            "Documento vazio: nenhum conteúdo extraível em '%s'" % path
        )

    # Cabeçalhos e rodapés — o Word Compare marca mudanças neles e um rodapé
    # removido (nº de versão, carimbo do escritório) é mudança real. Entram
    # como blocos ao FINAL (não perturbam o alinhamento do corpo) com
    # style_name sentinela p/ classificação METADATA e render rotulado.
    blocks.extend(_header_footer_blocks(docx_doc))

    for i, block in enumerate(blocks):
        block.index = i

    title = ""
    try:
        title = (docx_doc.core_properties.title or "").strip()
    except Exception:
        title = ""
    if not title:
        title = p.stem

    return Document(
        source_path=str(p),
        fmt="docx",
        blocks=blocks,
        page_count=0,
        title=title,
        default_font=default_font,
        default_font_size_pt=default_size_pt,
    )
