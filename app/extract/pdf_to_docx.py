"""Conversão PDF → DOCX (pdf2docx) para rotear pares PDF pelo pipeline Word.

Pares .pdf x .pdf historicamente caíam direto no gerador ReportLab
(app/output/redline_pdf.py), que re-tipografa o documento em layout
padronizado. Convertendo os dois PDFs para DOCX primeiro, o par passa a usar
o MESMO fluxo dos pares .docx: comparação canônica + redline in-place
preservando a formatação + PDF fiel via LibreOffice.

Só funciona para PDF nato-digital (com camada de texto). PDF escaneado
(imagem pura) geraria um DOCX vazio e um redline sem conteúdo — nesses casos
levantamos ``PdfConversionError`` e o chamador (app/jobs.py) volta ao gerador
padronizado, registrando warning no item.
"""
from __future__ import annotations

import copy
import logging
import os
import re

logger = logging.getLogger(__name__)

# Início de nova cláusula/item/título — nunca fundir o bloco seguinte no
# anterior quando o próximo começa assim (senão colaríamos duas cláusulas).
_CLAUSE_START_RE = re.compile(
    r"^\s*(\d+([.\-–][0-9A-Za-z]+)*[.\-–)]?\s"      # 1.  1.1.  10-A.  2)
    r"|\(?[a-z]\)"                                    # (a) a)
    r"|\(?[ivxIVX]+\)"                               # (i) (iv)
    r"|CLÁUSULA|PARÁGRAFO|Parágrafo|Art\.|Artigo"
    r"|CONSIDERANDO|CONSIDERANDOS|RESOLVEM|ANEXO)"
)
# Fim de frase/parágrafo: pontuação terminal (com aspas/parênteses de fecho).
_SENT_END_RE = re.compile(r"[.;:!?…][\"'”’)\]]?\s*$")

# pdf2docx loga no ROOT logger (logging.info direto, com basicConfig no
# import) — uma linha por página convertida. O filtro abaixo derruba esses
# registros apenas enquanto a conversão roda.
class _DropRootRecords(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        return record.name != "root"

# Mínimo de caracteres não-brancos no PDF inteiro para considerarmos que há
# camada de texto utilizável (abaixo disso tratamos como escaneado).
_MIN_TEXT_CHARS = 40


class PdfConversionError(ValueError):
    """PDF não pôde ser convertido para DOCX (escaneado, protegido, corrompido)."""


def _ensure_text_layer(pdf_path: str) -> None:
    import fitz  # PyMuPDF

    name = os.path.basename(pdf_path)
    try:
        doc = fitz.open(pdf_path)
    except Exception as exc:
        raise PdfConversionError("Não foi possível abrir '%s': %s" % (name, exc)) from exc
    try:
        if doc.needs_pass:
            raise PdfConversionError("PDF protegido por senha: '%s'." % name)
        chars = 0
        for page in doc:
            chars += len("".join(page.get_text().split()))
            if chars >= _MIN_TEXT_CHARS:
                return
    finally:
        doc.close()
    raise PdfConversionError(
        "PDF sem camada de texto (provavelmente escaneado): '%s'." % name
    )


def _tbl_col_count(tbl_el, qn) -> int:
    """Nº de colunas de um ``w:tbl`` (células da primeira linha)."""
    for tr in tbl_el.iter(qn("w:tr")):
        return len(tr.findall(qn("w:tc")))
    return 0


def _tr_text(tr, qn) -> str:
    return " | ".join(
        "".join(t.text or "" for t in tc.iter(qn("w:t"))).strip()
        for tc in tr.findall(qn("w:tc"))
    )


def _merge_adjacent_tables(docx_path: str) -> int:
    """Reúne tabelas que o pdf2docx PARTIU em várias (correção B6).

    Em PDF nativo o pdf2docx quebra uma única tabela em fragmentos: cabeçalho
    numa tabela de 1 linha + corpo em outra, ou as linhas em dois blocos
    (parcelas 1-2 | 3-5). Fragmentos com a MESMA quantidade de colunas,
    separados apenas por parágrafos vazios, são a mesma tabela — fundimos as
    linhas de volta. Roda ANTES de achatar pseudo-tabelas (senão o cabeçalho
    de 1 linha viraria parágrafo e apareceria marcado como alteração) e é
    determinístico, aplicado igual a base e revisado.

    Retorna quantos fragmentos foram fundidos. Nunca levanta.
    """
    from docx import Document
    from docx.oxml.ns import qn

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.warning("Não foi possível reabrir '%s' p/ fundir tabelas (%s).",
                       os.path.basename(docx_path), exc)
        return 0

    body = doc.element.body
    merged = 0
    last_tbl = None
    for child in list(body.iterchildren()):
        tag = child.tag
        if tag == qn("w:tbl"):
            if (
                last_tbl is not None
                and _tbl_col_count(last_tbl, qn) > 0
                and _tbl_col_count(child, qn) == _tbl_col_count(last_tbl, qn)
            ):
                header = _tr_text(next(iter(last_tbl.iter(qn("w:tr")))), qn)
                for tr in child.findall(qn("w:tr")):
                    # Não duplica um cabeçalho repetido no fragmento seguinte.
                    if _tr_text(tr, qn) == header:
                        continue
                    last_tbl.append(copy.deepcopy(tr))
                body.remove(child)
                merged += 1
                continue  # last_tbl continua sendo o alvo (funde 3+ fragmentos)
            last_tbl = child
        elif tag == qn("w:p"):
            if _para_text(child, qn).strip():
                last_tbl = None  # prosa real separa tabelas distintas
        else:
            last_tbl = None

    if merged:
        try:
            doc.save(docx_path)
            logger.info("Fragmentos de tabela fundidos em %s: %d",
                        os.path.basename(docx_path), merged)
        except Exception as exc:
            logger.warning("Falha ao salvar DOCX com tabelas fundidas '%s' (%s).",
                           os.path.basename(docx_path), exc)
            return 0
    return merged


def _is_pseudo_table(tbl) -> bool:
    """True se a tabela é, na verdade, um parágrafo corrido que o pdf2docx
    envolveu em grade (uma palavra por célula).

    Sinais: (a) uma única linha (qualquer largura), ou (b) muitas colunas
    (≥5) com a maioria das células de uma só palavra — parágrafo que quebrou
    em N linhas visuais vira tabela N×M. Tabelas REAIS destes contratos têm
    poucas colunas (2–4) com cabeçalho + dados, então não disparam (b).
    """
    n_rows = len(tbl.rows)
    if n_rows == 1:
        return True
    n_cols = len(tbl.columns)
    if n_cols < 5:
        return False
    cells = [c.text.strip() for row in tbl.rows for c in row.cells]
    nonempty = [txt for txt in cells if txt]
    if not nonempty:
        return False
    single_word = sum(1 for txt in nonempty if " " not in txt)
    return (single_word / len(nonempty)) >= 0.6


def _flatten_pseudo_tables(docx_path: str) -> int:
    """Achata pseudo-tabelas do pdf2docx em parágrafos (correção B3).

    Em PDF nativo, o pdf2docx frequentemente envolve parágrafos corridos em
    TABELAS com uma palavra por célula (ex.: "CONSIDERANDO | que | a |
    CONTRATANTE | …"). Renderizadas, quebram a visualização (vãos enormes,
    palavras partidas) e degradam o alinhamento do diff — e prendem trechos
    do texto que deveriam ser comparados como prosa (contribui para B1).
    Detectadas por ``_is_pseudo_table``, são desmontadas de volta a um
    parágrafo, na ordem de leitura (linha a linha), preservando os runs.

    Retorna quantas tabelas foram achatadas. Nunca levanta — falha aqui não
    deve derrubar a conversão.
    """
    from docx import Document
    from docx.oxml.ns import qn

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.warning("Não foi possível reabrir '%s' p/ achatar tabelas (%s).",
                       os.path.basename(docx_path), exc)
        return 0

    body = doc.element.body
    flattened = 0
    for tbl in list(doc.tables):
        tbl_el = tbl._element
        # Só tabelas de nível superior (evita mexer em tabelas aninhadas).
        if tbl_el.getparent() is not body:
            continue
        if not _is_pseudo_table(tbl):
            continue  # tabelas reais (grade de campos) permanecem intactas

        new_p = tbl_el.makeelement(qn("w:p"), {})
        ppr_copied = False
        last_text = ""
        # Ordem de leitura: todas as linhas, célula a célula (um parágrafo que
        # quebrou em N linhas visuais virou uma tabela de N linhas).
        cells_in_order = [c for r in tbl.rows for c in r.cells]
        prev_cell_text = None
        for cell in cells_in_order:
            cell_text = cell.text.strip()
            # O pdf2docx às vezes duplica palavras em células vizinhas
            # ("tiverem | tiverem"). Como só achatamos PROSA (pseudo-tabela),
            # célula idêntica à anterior é sempre artefato — descarta.
            if cell_text and cell_text == prev_cell_text:
                continue
            if cell_text:
                prev_cell_text = cell_text
            for para in cell.paragraphs:
                ppr = para._p.find(qn("w:pPr"))
                if not ppr_copied and ppr is not None:
                    ppr_copy = copy.deepcopy(ppr)
                    # Remove o recuo herdado da célula (posição x no PDF) — num
                    # parágrafo normal isso vira 1ª linha deslocada. Mantém
                    # alinhamento/espaçamento.
                    for ind in ppr_copy.findall(qn("w:ind")):
                        ppr_copy.remove(ind)
                    new_p.append(ppr_copy)
                    ppr_copied = True
                for run in para.runs:
                    new_p.append(copy.deepcopy(run._element))
                    last_text = run.text or last_text
            # Garante um espaço entre células vizinhas se ainda não houver.
            if last_text and not last_text.endswith((" ", "\t", " ")):
                sep = new_p.makeelement(qn("w:r"), {})
                t = sep.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
                t.text = " "
                sep.append(t)
                new_p.append(sep)
                last_text = " "

        tbl_el.addprevious(new_p)
        tbl_el.getparent().remove(tbl_el)
        flattened += 1

    if flattened:
        try:
            doc.save(docx_path)
            logger.info("Pseudo-tabelas achatadas em %s: %d",
                        os.path.basename(docx_path), flattened)
        except Exception as exc:
            logger.warning("Falha ao salvar DOCX achatado '%s' (%s).",
                           os.path.basename(docx_path), exc)
            return 0
    return flattened


def _para_text(p, qn) -> str:
    """Texto concatenado de um elemento ``w:p`` (todos os ``w:t``)."""
    parts = []
    for t in p.iter(qn("w:t")):
        parts.append(t.text or "")
    return "".join(parts)


def _is_heading_like(text: str) -> bool:
    """Título/rótulo curto: caixa-alta dominante e sem pontuação final, ou
    começando por 'CLÁUSULA'/'PARÁGRAFO'. Nunca é fundido com prosa vizinha."""
    stripped = text.strip()
    if not stripped:
        return False
    if _CLAUSE_START_RE.match(stripped) and not stripped[0].isdigit():
        return True  # CLÁUSULA / PARÁGRAFO / CONSIDERANDO …
    if len(stripped) <= 90 and not _SENT_END_RE.search(stripped):
        letters = [c for c in stripped if c.isalpha()]
        if letters and sum(1 for c in letters if c.isupper()) / len(letters) >= 0.7:
            return True
    return False


def _merge_two_paras(prev_p, cur_p, qn) -> None:
    """Move os runs de ``cur_p`` para o fim de ``prev_p``, tratando o espaço
    de junção (de-hifenização de quebra de linha ou espaço simples)."""
    prev_text = _para_text(prev_p, qn)
    cur_first = _para_text(cur_p, qn).lstrip()[:1]
    stripped = prev_text.rstrip()
    dehyphenate = (
        stripped.endswith("-")
        and len(stripped) >= 2
        and stripped[-2].isalpha()
        and cur_first.isalpha()
    )
    if dehyphenate:
        # Remove o hífen do último ``w:t`` não vazio de prev.
        for t in reversed(list(prev_p.iter(qn("w:t")))):
            if t.text and t.text.rstrip().endswith("-"):
                t.text = t.text.rstrip()[:-1]
                break
    elif prev_text and not prev_text.endswith((" ", "\t", "\n")):
        sep = prev_p.makeelement(qn("w:r"), {})
        st = sep.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
        st.text = " "
        sep.append(st)
        prev_p.append(sep)
    for r in cur_p.findall(qn("w:r")):
        prev_p.append(copy.deepcopy(r))


def _merge_wrapped_paragraphs(docx_path: str) -> int:
    """Refunde parágrafos que o pdf2docx quebrou em vários ``w:p`` (correção B1).

    Uma cláusula longa vira N parágrafos (um por linha visual do PDF). Como a
    quebra difere entre base e revisado, o alinhamento vê modify/insert/delete
    fantasmas e o Summary infla. Aqui fundimos um parágrafo no anterior quando:
    o anterior NÃO termina em pontuação final, nenhum dos dois é título/rótulo,
    e o próximo NÃO começa como nova cláusula/item. É determinístico (mesma
    regra nos dois lados) e preserva o conteúdo — nenhuma alteração some.

    Retorna quantos parágrafos foram fundidos. Nunca levanta.
    """
    from docx import Document
    from docx.oxml.ns import qn

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.warning("Não foi possível reabrir '%s' p/ fundir parágrafos (%s).",
                       os.path.basename(docx_path), exc)
        return 0

    body = doc.element.body
    merged = 0
    prev_p = None
    for child in list(body.iterchildren()):
        if child.tag != qn("w:p"):
            prev_p = None  # tabela/seção quebra a sequência de prosa
            continue
        cur_text = _para_text(child, qn)
        if not cur_text.strip():
            continue  # parágrafo vazio: ignora, não quebra a corrente
        if (
            prev_p is not None
            and not _SENT_END_RE.search(_para_text(prev_p, qn))
            and not _is_heading_like(_para_text(prev_p, qn))
            and not _is_heading_like(cur_text)
            and not _CLAUSE_START_RE.match(cur_text)
        ):
            _merge_two_paras(prev_p, child, qn)
            body.remove(child)
            merged += 1
            continue
        prev_p = child

    if merged:
        try:
            doc.save(docx_path)
            logger.info("Parágrafos refundidos em %s: %d",
                        os.path.basename(docx_path), merged)
        except Exception as exc:
            logger.warning("Falha ao salvar DOCX refundido '%s' (%s).",
                           os.path.basename(docx_path), exc)
            return 0
    return merged


def _fix_run_boundary_spaces(docx_path: str) -> int:
    """Reinsere espaços que o pdf2docx come na FRONTEIRA entre runs.

    O pdf2docx às vezes quebra o parágrafo em runs (mudança de fonte/estilo) e
    perde o espaço da junção: "…pagará à" + "CONTRATADA" vira "àCONTRATADA".
    Correção conservadora: só insere espaço quando o run anterior termina em
    letra MINÚSCULA e o próximo começa em letra MAIÚSCULA (início de nova
    palavra/nome próprio) — nunca no meio de palavra em CAIXA-ALTA (ex.:
    "CONTRA"+"TANTE" fica intacto porque "CONTRA" termina em maiúscula).

    Retorna quantos espaços foram inseridos. Nunca levanta.
    """
    from docx import Document

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.warning("Não foi possível reabrir '%s' p/ ajustar espaços (%s).",
                       os.path.basename(docx_path), exc)
        return 0

    fixed = 0

    def _fix_paragraph(p):
        nonlocal fixed
        runs = p.runs
        for i in range(len(runs) - 1):
            a = runs[i].text or ""
            b = runs[i + 1].text or ""
            if not a or not b:
                continue
            ca, cb = a[-1], b[0]
            if (
                ca.isalpha() and ca.islower()
                and cb.isalpha() and cb.isupper()
                and not a.endswith((" ", "\t"))
            ):
                runs[i + 1].text = " " + b
                fixed += 1

    for p in doc.paragraphs:
        _fix_paragraph(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _fix_paragraph(p)

    if fixed:
        try:
            doc.save(docx_path)
            logger.info("Espaços de fronteira de run corrigidos em %s: %d",
                        os.path.basename(docx_path), fixed)
        except Exception as exc:
            logger.warning("Falha ao salvar DOCX com espaços corrigidos '%s' (%s).",
                           os.path.basename(docx_path), exc)
            return 0
    return fixed


def _pdf_page_size_pt(pdf_path: str):
    """Retorna (width_pt, height_pt) da página mais larga do PDF."""
    import fitz

    doc = fitz.open(pdf_path)
    try:
        if doc.page_count <= 0:
            return None, None
        best_w = best_h = 0.0
        for pno in range(doc.page_count):
            rect = doc.load_page(pno).rect
            w, h = float(rect.width), float(rect.height)
            if w > best_w:
                best_w, best_h = w, h
        if best_w > 0 and best_h > 0:
            return best_w, best_h
    finally:
        doc.close()
    return None, None


def _apply_pdf_page_geometry(pdf_path: str, docx_path: str) -> None:
    """Copia largura/altura/orientação do PDF para as seções do DOCX.

    O pdf2docx costuma gravar A4 retrato mesmo quando o PDF é paisagem —
    o redline in-place + LibreOffice então corta o conteúdo. Forçamos a
    geometria real do PDF de origem.
    """
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Pt

    width_pt, height_pt = _pdf_page_size_pt(pdf_path)
    if not width_pt or not height_pt:
        return

    try:
        doc = Document(docx_path)
    except Exception as exc:
        logger.warning(
            "Não foi possível reabrir '%s' p/ aplicar geometria (%s).",
            os.path.basename(docx_path), exc,
        )
        return

    landscape = width_pt > height_pt
    try:
        for section in doc.sections:
            section.orientation = (
                WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
            )
            # python-docx troca width/height ao mudar orientation — redefine.
            section.page_width = Pt(width_pt)
            section.page_height = Pt(height_pt)
            # Margens um pouco menores em paisagem para caber tabelas largas.
            if landscape:
                for attr in ("left_margin", "right_margin"):
                    try:
                        current = getattr(section, attr)
                        if current and float(current.pt) > 54:  # > 0.75"
                            setattr(section, attr, Pt(36))  # 0.5"
                    except Exception:
                        pass
        doc.save(docx_path)
        logger.info(
            "Geometria PDF→DOCX: %.0f×%.0f pt (%s) em %s",
            width_pt, height_pt,
            "paisagem" if landscape else "retrato",
            os.path.basename(docx_path),
        )
    except Exception as exc:
        logger.warning(
            "Falha ao aplicar geometria de página em '%s' (%s).",
            os.path.basename(docx_path), exc,
        )


def convert_pdf_to_docx(pdf_path: str, docx_path: str) -> None:
    """Converte um PDF nato-digital em DOCX editável.

    Levanta ``PdfConversionError`` quando o PDF não tem texto extraível ou a
    conversão falha — o chamador decide o fallback.
    """
    if not pdf_path or not os.path.isfile(pdf_path):
        raise PdfConversionError("Arquivo PDF não encontrado: '%s'" % pdf_path)

    _ensure_text_layer(pdf_path)

    out_dir = os.path.dirname(os.path.abspath(docx_path))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    try:
        from pdf2docx import Converter
    except Exception as exc:
        raise PdfConversionError(
            "Biblioteca pdf2docx indisponível: %s" % exc
        ) from exc

    converter = None
    quiet = _DropRootRecords()
    logging.getLogger().addFilter(quiet)
    try:
        converter = Converter(pdf_path)
        converter.convert(docx_path)
    except PdfConversionError:
        raise
    except Exception as exc:
        raise PdfConversionError(
            "Falha ao converter '%s' para DOCX: %s" % (os.path.basename(pdf_path), exc)
        ) from exc
    finally:
        logging.getLogger().removeFilter(quiet)
        if converter is not None:
            try:
                converter.close()
            except Exception:
                pass

    if not os.path.isfile(docx_path) or os.path.getsize(docx_path) == 0:
        raise PdfConversionError(
            "Conversão de '%s' não produziu DOCX válido." % os.path.basename(pdf_path)
        )
    # Correção B6: reúne tabelas partidas (cabeçalho/linhas em fragmentos)
    # ANTES de achatar pseudo-tabelas — assim um cabeçalho de 1 linha não é
    # achatado em parágrafo e marcado como alteração fantasma.
    _merge_adjacent_tables(docx_path)
    # Correção B3: desfaz as pseudo-tabelas (parágrafos corridos que o pdf2docx
    # envolve em tabela) antes de o par seguir para extração/comparação.
    _flatten_pseudo_tables(docx_path)
    # Correção B1: refunde parágrafos que o pdf2docx quebrou por linha visual
    # (fragmentação difere entre base/revisado → contagem fantasma no Summary).
    _merge_wrapped_paragraphs(docx_path)
    # Correção B4: reinsere espaços comidos na fronteira entre runs
    # ("pagará à" + "CONTRATADA" → "pagará àCONTRATADA").
    _fix_run_boundary_spaces(docx_path)
    # Paisagem / tamanho real da página do PDF (pdf2docx tende a forçar A4).
    _apply_pdf_page_geometry(pdf_path, docx_path)
    logger.info("PDF convertido para DOCX: %s -> %s", pdf_path, docx_path)
