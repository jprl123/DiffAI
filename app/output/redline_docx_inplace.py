"""Redline DOCX in-place — preserva formatação do documento revisado.

Copia o arquivo revisado e aplica apenas as marcações de alteração nos
parágrafos e tabelas existentes (estilos, margens, fontes do template
permanecem). Parágrafos exclusivos da base são inseridos no fluxo com
marcação de exclusão.
"""
from __future__ import annotations

import logging
import os
import shutil
from typing import List, Optional, Tuple, Union

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.shared import RGBColor

from app.extract.docx_extractor import (
    _image_hashes,
    _runs_from_paragraph,
)
from app.models import BlockKind, ChangeType, ComparisonResult, Fragment, RenderBlock
from app.output.summary import APP_NAME, SUMMARY_TITLE, format_compared_at, summary_rows

logger = logging.getLogger(__name__)

COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)
_ROW_RULE_COLOR = "D1D5DB"  # filete horizontal fino entre linhas da síntese

COLOR_INSERT = RGBColor(0x1A, 0x56, 0xDB)
COLOR_DELETE = RGBColor(0xDC, 0x26, 0x26)
COLOR_MOVE = RGBColor(0x0E, 0x7A, 0x3D)

_MOVE_TYPES = (ChangeType.MOVE, ChangeType.MOVE_MODIFY)
BodySlot = Tuple[str, Union[DocxParagraph, DocxTable]]


def _is_move(change_type: Optional[ChangeType]) -> bool:
    return change_type in _MOVE_TYPES


def _effective_op(frag_op: str, change_type: Optional[ChangeType]) -> str:
    op = (frag_op or "equal").lower()
    if op == "equal":
        if change_type == ChangeType.INSERT:
            return "insert"
        if change_type == ChangeType.DELETE:
            return "delete"
    return op


def _clear_runs(paragraph: DocxParagraph) -> None:
    element = paragraph._element
    for run in list(element.findall(qn("w:r"))):
        element.remove(run)


def _add_marked_run(paragraph: DocxParagraph, frag: Fragment, change_type: Optional[ChangeType]):
    text = frag.text if frag.text is not None else ""
    if not text:
        return
    run = paragraph.add_run(text)
    run.bold = bool(frag.bold)
    run.italic = bool(frag.italic)
    run.underline = bool(frag.underline)
    if frag.strike:
        run.font.strike = True

    op = _effective_op(getattr(frag, "op", "equal"), change_type)
    if op == "insert":
        run.font.color.rgb = COLOR_INSERT
        run.underline = True
    elif op == "delete":
        run.font.color.rgb = COLOR_DELETE
        run.font.strike = True
    elif _is_move(change_type):
        run.font.color.rgb = COLOR_MOVE


def _add_moved_suffix(paragraph: DocxParagraph) -> None:
    run = paragraph.add_run(" [movido]")
    run.font.color.rgb = COLOR_MOVE
    run.italic = True


def _collect_body_slots(doc: DocxDocument) -> List[BodySlot]:
    """Mesma ordem de blocos que ``extract_docx`` (texto/imagem/tabela)."""
    slots: List[BodySlot] = []
    part = doc.part
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph = DocxParagraph(child, doc)
            runs = _runs_from_paragraph(paragraph)
            if any(r.text.strip() for r in runs):
                slots.append(("text", paragraph))
            for _hash in _image_hashes(paragraph, part):
                slots.append(("image", paragraph))
        elif child.tag == qn("w:tbl"):
            slots.append(("table", DocxTable(child, doc)))
    return slots


def _insert_paragraph_before(
    doc: DocxDocument, ref: Optional[DocxParagraph], rb: RenderBlock
) -> None:
    """Insere parágrafo de exclusão antes do elemento de referência."""
    from docx.oxml import OxmlElement

    new_p = OxmlElement("w:p")
    body = doc.element.body
    if ref is not None:
        body.insert(body.index(ref._element), new_p)
    else:
        body.append(new_p)
    paragraph = DocxParagraph(new_p, doc)
    for frag in rb.fragments or []:
        _add_marked_run(paragraph, frag, ChangeType.DELETE)
    if _is_move(rb.change_type):
        _add_moved_suffix(paragraph)


def _mark_text_paragraph(paragraph: DocxParagraph, rb: RenderBlock) -> None:
    _clear_runs(paragraph)
    label_pending = rb.list_label  # rótulo novo injetado no diff (renumeração)
    for frag in rb.fragments or []:
        if frag is None:
            continue
        # O Word já renderiza a numeração automática nova ("(b)") neste
        # parágrafo — o fragmento INSERT do rótulo seria duplicado; apara.
        if (
            label_pending
            and frag.op == "insert"
            and frag.text.lstrip().startswith(label_pending)
        ):
            trimmed = frag.text.lstrip()[len(label_pending):].lstrip("\t ")
            label_pending = None
            if not trimmed:
                continue
            frag = Fragment(
                text=trimmed, op=frag.op, bold=frag.bold, italic=frag.italic,
                underline=frag.underline, strike=frag.strike,
            )
        _add_marked_run(paragraph, frag, rb.change_type)
    if _is_move(rb.change_type):
        _add_moved_suffix(paragraph)


def _insert_deleted_row(table: DocxTable, at_physical: int, row_frags) -> None:
    """Insere uma linha FÍSICA nova na tabela para exibir uma linha excluída.

    A tabela do documento revisado não contém as linhas removidas — sem esta
    inserção, exclusões de linha ficariam INVISÍVEIS no redline fiel (defeito
    real encontrado na tabela de definições do memorandum; a pior classe de
    erro possível num redline).
    """
    from docx.oxml import OxmlElement

    new_row = table.add_row()  # nasce no fim; movemos para a posição certa
    new_tr = new_row._tr
    if at_physical < len(table.rows) - 1:
        ref_tr = table.rows[at_physical]._tr
        ref_tr.addprevious(new_tr)
    for j, cell_frags in enumerate(row_frags):
        if j >= len(new_row.cells):
            break
        paragraph = new_row.cells[j].paragraphs[0]
        _clear_runs(paragraph)
        for frag in cell_frags or []:
            if frag is None:
                continue
            _add_marked_run(paragraph, frag, ChangeType.DELETE)
            run = paragraph.runs[-1]
            run.font.strike = True
            run.font.color.rgb = COLOR_DELETE


def _mark_table(table: DocxTable, rb: RenderBlock) -> None:
    """Marca a tabela linha a linha.

    ``rb.rows`` está na ordem do render (linhas excluídas INTERCALADAS);
    a tabela física só tem as linhas do documento revisado. O ponteiro
    ``physical`` avança apenas em linhas que existem de fato — linhas
    excluídas são inseridas fisicamente na posição corrente.
    """
    rows = rb.rows or []
    row_ops = rb.row_ops or []
    physical = 0
    for i, row_frags in enumerate(rows):
        row_op = ((row_ops[i] if i < len(row_ops) else "equal") or "equal").lower()
        if row_op == "delete":
            _insert_deleted_row(table, physical, row_frags)
            physical += 1  # a linha inserida agora ocupa a posição corrente
            continue
        if physical >= len(table.rows):
            logger.warning("Redline in-place: linhas render além da tabela física")
            break
        for j, cell_frags in enumerate(row_frags):
            if j >= len(table.rows[physical].cells):
                break
            cell = table.rows[physical].cells[j]
            paragraph = cell.paragraphs[0]
            _clear_runs(paragraph)
            for frag in cell_frags or []:
                if frag is None:
                    continue
                _add_marked_run(paragraph, frag, rb.change_type)
                run = paragraph.runs[-1]
                if row_op == "insert":
                    run.font.color.rgb = COLOR_INSERT
                    run.underline = True
        physical += 1


def _apply_render_blocks(doc: DocxDocument, render_blocks: List[RenderBlock]) -> None:
    slots = _collect_body_slots(doc)
    slot_idx = 0

    for rb in render_blocks:
        if rb is None:
            continue
        # Cabeçalho/rodapé não têm slot no corpo do DOCX — as mudanças deles
        # aparecem no PDF padronizado e no relatório analítico (METADATA).
        if rb.section_path and rb.section_path[0] in ("Cabeçalho", "Rodapé"):
            continue
        if rb.change_type == ChangeType.DELETE:
            ref = None
            if slot_idx < len(slots) and slots[slot_idx][0] == "text":
                ref = slots[slot_idx][1]
            elif slot_idx > 0 and slots[slot_idx - 1][0] == "text":
                ref = slots[slot_idx - 1][1]
            if rb.kind == BlockKind.TABLE:
                # Tabela removida: inserir marcador textual antes da posição atual
                _insert_paragraph_before(
                    doc, ref if isinstance(ref, DocxParagraph) else None, rb
                )
            else:
                _insert_paragraph_before(
                    doc, ref if isinstance(ref, DocxParagraph) else None, rb
                )
            continue

        if slot_idx >= len(slots):
            logger.warning("Redline in-place: blocos render > slots do documento")
            break

        kind, obj = slots[slot_idx]
        if rb.kind == BlockKind.TABLE and kind == "table":
            _mark_table(obj, rb)  # type: ignore[arg-type]
            slot_idx += 1
        elif rb.kind in (BlockKind.PARAGRAPH, BlockKind.HEADING, BlockKind.LIST_ITEM) and kind == "text":
            _mark_text_paragraph(obj, rb)  # type: ignore[arg-type]
            slot_idx += 1
        elif rb.kind == BlockKind.IMAGE and kind == "image":
            # Imagem: marca o parágrafo âncora com nota curta se alterada
            if rb.change_type != ChangeType.EQUAL:
                para = obj  # type: ignore[assignment]
                _clear_runs(para)
                label = "[Imagem alterada]"
                if rb.change_type == ChangeType.INSERT:
                    label = "[Imagem inserida]"
                elif rb.change_type == ChangeType.DELETE:
                    label = "[Imagem removida]"
                run = para.add_run(label)
                run.italic = True
                run.font.color.rgb = COLOR_INSERT if rb.change_type == ChangeType.INSERT else COLOR_DELETE
            slot_idx += 1
        else:
            # Desalinhamento leve — avança para não travar o restante
            slot_idx += 1


def _set_cell_bottom_rule(cell) -> None:
    """Aplica um filete inferior fino à célula (visual limpo, sem grade)."""
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), _ROW_RULE_COLOR)
    borders.append(bottom)


def _append_summary_page(doc: DocxDocument, result: ComparisonResult) -> None:
    """Página de síntese ao FINAL do documento (última página do PDF fiel)."""
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    from app.output.summary import app_logo_path

    # Quebra de SEÇÃO (não de página): a síntese ganha cabeçalho e rodapé
    # próprios, vazios — sem herdar "Exhibit", numeração ou marca do documento.
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    for part in (
        section.header,
        section.footer,
        section.even_page_header,
        section.even_page_footer,
        section.first_page_header,
        section.first_page_footer,
    ):
        try:
            part.is_linked_to_previous = False
        except (AttributeError, ValueError):
            pass

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = brand.add_run(APP_NAME)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_MUTED

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SUMMARY_TITLE)
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(format_compared_at(result.compared_at))
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()

    rows = summary_rows(result)
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[i].cells
        label_cell.width = Cm(7.5)
        value_cell.width = Cm(8.5)
        p = label_cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        p = value_cell.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(10)
        _set_cell_bottom_rule(label_cell)
        _set_cell_bottom_rule(value_cell)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Gerado por %s" % APP_NAME)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    logo = app_logo_path()
    if logo:
        doc.add_paragraph()
        end_logo = doc.add_paragraph()
        end_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            end_logo.add_run().add_picture(logo, width=Cm(1.0))
        except (OSError, ValueError):
            pass


def write_redline_docx_inplace(
    result: ComparisonResult,
    compare_path: str,
    out_path: str,
) -> None:
    """Gera redline editável preservando o layout do documento revisado."""
    if result is None or not compare_path or not str(compare_path).strip():
        raise ValueError("Parâmetros inválidos para redline DOCX in-place.")
    if not os.path.isfile(compare_path):
        raise ValueError("Arquivo revisado não encontrado: '%s'" % compare_path)

    out_path = str(out_path)
    out_dir = os.path.dirname(os.path.abspath(out_path))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)

    shutil.copy2(compare_path, out_path)
    doc = DocxDocument(out_path)
    blocks = result.render_blocks or []
    if blocks:
        _apply_render_blocks(doc, blocks)
    _append_summary_page(doc, result)
    try:
        doc.save(out_path)
    except Exception as exc:
        raise ValueError(
            "Falha ao gravar o DOCX de redline in-place em '%s': %s" % (out_path, exc)
        ) from exc
    logger.info("Redline DOCX in-place gravado em %s", out_path)
