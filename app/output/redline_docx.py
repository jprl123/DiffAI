"""Saída DOCX editável do redline (python-docx).

Percorre ``result.render_blocks`` e produz um documento Word com as
alterações marcadas: inserções em azul sublinhado, exclusões em vermelho
tachado, movimentações em verde com sufixo " [movido]". Tabelas são
reproduzidas célula a célula com a mesma marcação. A primeira página traz
título e legenda; a última, a página de síntese ("Summary of Changes").
"""
from __future__ import annotations

import logging
import os
from typing import List, Optional

from docx import Document as _DocxDocument
from docx.enum.text import WD_BREAK
from docx.shared import Pt, RGBColor

from app.models import (
    BlockKind,
    ChangeType,
    ComparisonResult,
    Fragment,
    RenderBlock,
)

logger = logging.getLogger(__name__)

# Cores de marcação (mesma paleta do redline PDF)
COLOR_INSERT = RGBColor(0x1A, 0x56, 0xDB)   # azul
COLOR_DELETE = RGBColor(0xDC, 0x26, 0x26)   # vermelho
COLOR_MOVE = RGBColor(0x0E, 0x7A, 0x3D)     # verde

_MOVE_TYPES = (ChangeType.MOVE, ChangeType.MOVE_MODIFY)


# ---------------------------------------------------------------------------
# Helpers internos
# ---------------------------------------------------------------------------

def _is_move(change_type: Optional[ChangeType]) -> bool:
    return change_type in _MOVE_TYPES


def _effective_op(frag_op: str, change_type: Optional[ChangeType]) -> str:
    """Resolve o op efetivo do fragmento considerando o tipo do bloco.

    Em blocos INSERT/DELETE inteiros, fragmentos "equal" herdam a marcação
    do bloco (defensivo — o motor normalmente já preenche o op correto).
    """
    op = (frag_op or "equal").lower()
    if op == "equal":
        if change_type == ChangeType.INSERT:
            return "insert"
        if change_type == ChangeType.DELETE:
            return "delete"
    return op


def _add_fragment_run(paragraph, frag: Fragment, change_type: Optional[ChangeType]):
    """Adiciona um run ao parágrafo com a formatação do fragmento + marcação."""
    text = frag.text if frag.text is not None else ""
    run = paragraph.add_run(text)
    # Formatação original do fragmento
    run.bold = bool(frag.bold)
    run.italic = bool(frag.italic)
    run.underline = bool(frag.underline)
    if frag.strike:
        run.font.strike = True

    op = _effective_op(getattr(frag, "op", "equal"), change_type)
    if op == "insert":
        run.font.color.rgb = COLOR_INSERT
        run.underline = True
    elif op == "delete":
        run.font.color.rgb = COLOR_DELETE
        run.font.strike = True
    elif _is_move(change_type):
        # fragmentos "equal" de blocos movidos ficam verdes
        run.font.color.rgb = COLOR_MOVE
    return run


def _add_moved_suffix(paragraph) -> None:
    run = paragraph.add_run(" [movido]")
    run.font.color.rgb = COLOR_MOVE
    run.italic = True


def _clamp_heading_level(level: int) -> int:
    try:
        lvl = int(level)
    except (TypeError, ValueError):
        lvl = 1
    if lvl < 1:
        lvl = 1
    if lvl > 9:
        lvl = 9
    return lvl


def _render_text_block(doc, block: RenderBlock) -> None:
    """HEADING / PARAGRAPH / LIST_ITEM -> parágrafo com runs marcados."""
    if block.kind == BlockKind.HEADING:
        paragraph = doc.add_heading("", level=_clamp_heading_level(block.level))
    elif block.kind == BlockKind.LIST_ITEM:
        try:
            paragraph = doc.add_paragraph(style="List Bullet")
        except KeyError:
            paragraph = doc.add_paragraph()
            paragraph.add_run("• ")
    else:
        paragraph = doc.add_paragraph()

    fragments = block.fragments or []
    for frag in fragments:
        if frag is None:
            continue
        _add_fragment_run(paragraph, frag, block.change_type)

    if _is_move(block.change_type):
        _add_moved_suffix(paragraph)


def _render_image_block(doc, block: RenderBlock) -> None:
    """Bloco de imagem: marcador textual coerente com o tipo da mudança."""
    labels = {
        ChangeType.INSERT: ("[Imagem inserida]", COLOR_INSERT),
        ChangeType.DELETE: ("[Imagem removida]", COLOR_DELETE),
        ChangeType.MODIFY: ("[Imagem substituída]", COLOR_INSERT),
        ChangeType.MOVE: ("[Imagem movida]", COLOR_MOVE),
        ChangeType.MOVE_MODIFY: ("[Imagem movida e substituída]", COLOR_MOVE),
    }
    text, color = labels.get(block.change_type, ("[Imagem]", None))
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    if color is not None:
        run.font.color.rgb = color
    if block.change_type == ChangeType.DELETE:
        run.font.strike = True
    if _is_move(block.change_type):
        _add_moved_suffix(paragraph)


def _render_table_block(doc, block: RenderBlock) -> None:
    """Tabela: células preenchidas run a run com as mesmas cores.

    Linha com op "delete": todos os runs tachados em vermelho.
    Linha com op "insert": todos os runs em azul sublinhado.
    """
    rows = block.rows or []
    if not rows:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("[tabela sem conteúdo]")
        run.italic = True
        return

    n_rows = len(rows)
    n_cols = max(len(row) for row in rows) if rows else 0
    if n_cols == 0:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("[tabela sem colunas]")
        run.italic = True
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    row_ops = block.row_ops or []
    for i, row in enumerate(rows):
        row_op = (row_ops[i] if i < len(row_ops) else "equal") or "equal"
        row_op = row_op.lower()
        for j in range(n_cols):
            fragments = row[j] if j < len(row) else []
            cell = table.cell(i, j)
            paragraph = cell.paragraphs[0]
            for frag in fragments or []:
                if frag is None:
                    continue
                run = _add_fragment_run(paragraph, frag, block.change_type)
                if row_op == "delete":
                    run.font.strike = True
                    run.font.color.rgb = COLOR_DELETE
                elif row_op == "insert":
                    run.font.color.rgb = COLOR_INSERT
                    run.underline = True

    if _is_move(block.change_type):
        _add_moved_suffix(doc.add_paragraph())


def _add_page_break(doc) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _fmt_pages(pages: List[int]) -> str:
    if not pages:
        return "—"
    return ", ".join(str(p) for p in pages)


def _summary_items(result: ComparisonResult) -> List[tuple]:
    """Mesmos campos da página de síntese do PDF."""
    stats = result.stats
    base_name = os.path.basename(result.base_path) or result.base_title or "—"
    compare_name = (
        os.path.basename(result.compare_path) or result.compare_title or "—"
    )
    duration = getattr(result, "duration_seconds", 0.0) or 0.0
    total = int(stats.insertions) + int(stats.deletions) + int(stats.moves)
    return [
        ("Data da comparação", result.compared_at or "—"),
        ("Arquivo base", base_name),
        ("Arquivo revisado", compare_name),
        ("Total de alterações", str(total)),
        ("Inserções", str(stats.insertions)),
        ("Exclusões", str(stats.deletions)),
        ("Movimentações", str(stats.moves)),
        ("Modificações (in-place)", str(stats.modifications)),
        ("Mudanças de conteúdo", str(stats.content_changes)),
        ("Mudanças de formatação", str(stats.formatting_changes)),
        ("Mudanças rotineiras (ruído)", str(stats.noise_changes)),
        ("Mudanças em tabelas", str(stats.table_changes)),
        ("Mudanças em imagens", str(stats.image_changes)),
        ("Páginas alteradas", _fmt_pages(stats.changed_pages)),
        ("Duração (s)", "%.2f" % duration),
    ]


def _add_summary_page(doc, result: ComparisonResult) -> None:
    _add_page_break(doc)
    doc.add_heading("Summary of Changes", level=1)
    items = _summary_items(result)
    table = doc.add_table(rows=len(items), cols=2)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for i, (metric, value) in enumerate(items):
        cell_metric = table.cell(i, 0)
        run = cell_metric.paragraphs[0].add_run(metric)
        run.bold = True
        table.cell(i, 1).paragraphs[0].add_run(value)


def _add_cover(doc, result: ComparisonResult) -> None:
    base_name = os.path.basename(result.base_path) or result.base_title or "base"
    compare_name = (
        os.path.basename(result.compare_path) or result.compare_title or "revisado"
    )
    doc.add_heading("Redline: %s vs %s" % (base_name, compare_name), level=0)

    if result.compared_at:
        info = doc.add_paragraph()
        run = info.add_run("Comparação gerada em %s." % result.compared_at)
        run.italic = True
        run.font.size = Pt(9)

    legend = doc.add_paragraph()
    legend.add_run("Legenda: ")
    run = legend.add_run("texto inserido")
    run.font.color.rgb = COLOR_INSERT
    run.underline = True
    legend.add_run("   ")
    run = legend.add_run("texto excluído")
    run.font.color.rgb = COLOR_DELETE
    run.font.strike = True
    legend.add_run("   ")
    run = legend.add_run("texto movido [movido]")
    run.font.color.rgb = COLOR_MOVE
    legend.add_run("   (mudanças rotineiras aparecem com a mesma marcação "
                   "e são consolidadas na síntese).")


# ---------------------------------------------------------------------------
# API pública
# ---------------------------------------------------------------------------

def write_redline_docx(result: ComparisonResult, out_path: str) -> None:
    """Gera o DOCX editável do redline a partir de um ComparisonResult.

    Levanta ValueError (mensagem em pt-BR) para entradas inválidas ou falha
    de gravação.
    """
    if result is None:
        raise ValueError("Resultado de comparação inválido: recebido None.")
    render_blocks = getattr(result, "render_blocks", None)
    if render_blocks is None or getattr(result, "stats", None) is None:
        raise ValueError(
            "Resultado de comparação inválido: estrutura ComparisonResult "
            "incompleta (render_blocks/stats ausentes)."
        )
    if not out_path or not str(out_path).strip():
        raise ValueError("Caminho de saída do DOCX não informado.")

    out_path = str(out_path)
    out_dir = os.path.dirname(os.path.abspath(out_path))
    try:
        os.makedirs(out_dir, exist_ok=True)
    except OSError as exc:
        raise ValueError(
            "Não foi possível criar o diretório de saída '%s': %s"
            % (out_dir, exc)
        )

    doc = _DocxDocument()
    _add_cover(doc, result)

    if not render_blocks:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(
            "Nenhum bloco a renderizar — os documentos são idênticos "
            "ou estão vazios."
        )
        run.italic = True
    else:
        for block in render_blocks:
            if block is None:
                continue
            try:
                if block.kind == BlockKind.TABLE:
                    _render_table_block(doc, block)
                elif block.kind == BlockKind.IMAGE:
                    _render_image_block(doc, block)
                else:
                    _render_text_block(doc, block)
            except (AttributeError, TypeError, IndexError) as exc:
                logger.warning("Bloco de render inválido ignorado: %s", exc)
                continue

    _add_summary_page(doc, result)

    try:
        doc.save(out_path)
    except Exception as exc:
        raise ValueError(
            "Falha ao gravar o DOCX de redline em '%s': %s" % (out_path, exc)
        )
    logger.info("Redline DOCX gravado em %s", out_path)
