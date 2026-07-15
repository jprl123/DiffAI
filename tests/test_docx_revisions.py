"""Testes do aceite automático de track changes pendentes (pré-compare)."""
from __future__ import annotations

import os
import shutil
import tempfile
import time
import unittest

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.extract.docx_revisions import accept_all_revisions, has_tracked_revisions


def _el(parent, tag: str, attrs=None, text=None):
    element = parent.makeelement(qn(tag), {})
    for k, v in (attrs or {}).items():
        element.set(qn(k), v)
    if text is not None:
        element.text = text
    return element


_REV_ATTRS = {"w:id": "1", "w:author": "Teste", "w:date": "2026-01-01T00:00:00Z"}


def _make_revisioned_docx(path: str) -> None:
    """Parágrafo com inserção E exclusão pendentes:
    'Cláusula ' + ins('nova ') + del('antiga ') + 'vigente.'
    Aceito → 'Cláusula nova vigente.'"""
    doc = DocxDocument()
    p = doc.add_paragraph("Cláusula ")

    ins = _el(p._p, "w:ins", _REV_ATTRS)
    run = _el(ins, "w:r")
    text = _el(run, "w:t", {"xml:space": "preserve"}, "nova ")
    run.append(text)
    ins.append(run)
    p._p.append(ins)

    dele = _el(p._p, "w:del", _REV_ATTRS)
    drun = _el(dele, "w:r")
    dtext = _el(drun, "w:delText", {"xml:space": "preserve"}, "antiga ")
    drun.append(dtext)
    dele.append(drun)
    p._p.append(dele)

    p.add_run("vigente.")
    doc.save(path)


def _make_deleted_mark_docx(path: str) -> None:
    """Marca do 1º parágrafo deletada → funde com o 2º ao aceitar."""
    doc = DocxDocument()
    p1 = doc.add_paragraph("Primeira parte ")
    doc.add_paragraph("segunda parte.")
    ppr = p1._p.get_or_add_pPr()
    rpr = _el(ppr, "w:rPr")
    rpr.append(_el(rpr, "w:del", _REV_ATTRS))
    ppr.append(rpr)
    doc.save(path)


class DocxRevisionsUnitTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.mkdtemp(prefix="comparedocs-rev-")

    def tearDown(self) -> None:
        shutil.rmtree(self.tmp, ignore_errors=True)

    def test_detection(self) -> None:
        clean = os.path.join(self.tmp, "limpo.docx")
        doc = DocxDocument()
        doc.add_paragraph("Sem revisões, mas com tabela.")
        doc.add_table(rows=1, cols=1)  # <w:insideH> não é falso positivo
        doc.save(clean)
        self.assertFalse(has_tracked_revisions(clean))

        dirty = os.path.join(self.tmp, "com_revisoes.docx")
        _make_revisioned_docx(dirty)
        self.assertTrue(has_tracked_revisions(dirty))

    def test_accept_insert_and_delete(self) -> None:
        path = os.path.join(self.tmp, "doc.docx")
        _make_revisioned_docx(path)
        accept_all_revisions(path)

        self.assertFalse(has_tracked_revisions(path))
        doc = DocxDocument(path)
        texts = [p.text for p in doc.paragraphs if p.text]
        self.assertEqual(texts, ["Cláusula nova vigente."])

    def test_accept_deleted_paragraph_mark_merges(self) -> None:
        path = os.path.join(self.tmp, "merge.docx")
        _make_deleted_mark_docx(path)
        accept_all_revisions(path)

        doc = DocxDocument(path)
        texts = [p.text for p in doc.paragraphs if p.text]
        self.assertEqual(texts, ["Primeira parte segunda parte."])

    def test_job_accepts_revisions_before_compare(self) -> None:
        """Par com revisões pendentes no revisado: compara o texto ACEITO,
        avisa no item e não toca no arquivo original."""
        from app.history import HistoryStore
        from app.jobs import JobManager

        base_path = os.path.join(self.tmp, "base.docx")
        doc = DocxDocument()
        doc.add_paragraph("Cláusula antiga vigente.")
        doc.save(base_path)

        compare_path = os.path.join(self.tmp, "revisado.docx")
        _make_revisioned_docx(compare_path)  # aceito = "Cláusula nova vigente."

        out_dir = os.path.join(self.tmp, "out")
        manager = JobManager(
            history_store=HistoryStore(path=os.path.join(self.tmp, "h.json"))
        )
        job_id = manager.create_job([(base_path, compare_path)], {"output_dir": out_dir})
        deadline = time.time() + 120
        while manager.get_job(job_id)["status"] == "running":
            self.assertLess(time.time(), deadline, "job não terminou a tempo")
            time.sleep(0.2)

        job = manager.get_job(job_id)
        self.assertEqual(job["status"], "done")
        item = job["items"][0]
        self.assertEqual(item["status"], "ok")
        self.assertTrue(
            any("marcas de revisão pendentes" in w for w in item["warnings"]),
            "aviso de revisões aceitas ausente: %r" % item["warnings"],
        )

        # A comparação enxergou o texto ACEITO (antiga → nova).
        result = manager.get_result_object(job_id, 0)
        old_new = [
            ((c.old_text or ""), (c.new_text or "")) for c in result.changes
        ]
        self.assertTrue(
            any("nova" in new for _old, new in old_new),
            "inserção pendente não entrou na comparação: %r" % old_new,
        )

        # Original do usuário intocado; redline de saída sem marcas antigas.
        self.assertTrue(has_tracked_revisions(compare_path))
        self.assertIn("docx", item["outputs"])
        self.assertFalse(has_tracked_revisions(item["outputs"]["docx"]))


if __name__ == "__main__":
    unittest.main()
