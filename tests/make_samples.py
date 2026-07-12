"""Geração de amostras realistas (pt-BR) para os testes ponta-a-ponta do Compare-Docs.

Cria tests/samples/{base,revised}/ com três pares de documentos:

1. "Contrato de Prestação de Serviços.docx" — contrato de ~3 páginas com título,
   seções numeradas (Heading 1/2), ~20 parágrafos jurídicos, tabela de valores e
   linha de versão. A versão revisada contém um conjunto CONHECIDO de mudanças
   (versão/data, prazo 30->45, multa 2%->10%, parágrafo novo, parágrafo removido,
   cláusula de foro movida para o final, célula de tabela alterada, linha nova na
   tabela e um parágrafo com apenas formatação alterada).
2. "Política de Privacidade.pdf" — PDF de 2-3 páginas gerado com PyMuPDF (fitz);
   revisado com data atualizada, um parágrafo alterado e um parágrafo inserido.
3. "Proposta Comercial v1.docx" (base) / "Proposta Comercial v2 final.docx"
   (revisado) — nomes diferentes para exercitar o pareamento fuzzy; 2 mudanças.
4. "Orçamento Projeto.xlsx" — planilha com abas Custos e Resumo; células e
   linhas alteradas na revisão.

Idempotente: apaga e recria tests/samples/ a cada execução.

Documentos reais (ex.: Rivio) ficam em tests/fixtures/ e NÃO são tocados aqui.

Uso:
    .venv/bin/python -m tests.make_samples
"""
from __future__ import annotations

import os
import shutil
from typing import List, Tuple

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import Workbook

TESTS_DIR = os.path.dirname(os.path.abspath(__file__))
SAMPLES_DIR = os.path.join(TESTS_DIR, "samples")
BASE_DIR = os.path.join(SAMPLES_DIR, "base")
REVISED_DIR = os.path.join(SAMPLES_DIR, "revised")

CONTRACT_NAME = "Contrato de Prestação de Serviços.docx"
POLICY_NAME = "Política de Privacidade.pdf"
PROPOSAL_BASE_NAME = "Proposta Comercial v1.docx"
PROPOSAL_REVISED_NAME = "Proposta Comercial v2 final.docx"
BUDGET_NAME = "Orçamento Projeto.xlsx"


# ---------------------------------------------------------------------------
# Conteúdo do contrato (constantes compartilhadas entre base e revisado)
# ---------------------------------------------------------------------------

CONTRACT_TITLE = "Contrato de Prestação de Serviços"
VERSION_LINE_BASE = "Versão 1.0 — 10/01/2026"
VERSION_LINE_REVISED = "Versão 2.0 — 05/07/2026"

INTRO_1 = (
    "Pelo presente instrumento particular, de um lado EMPRESA ALFA CONSULTORIA LTDA., "
    "pessoa jurídica de direito privado, inscrita no CNPJ sob o nº 12.345.678/0001-90, "
    "com sede na Avenida Paulista, nº 1.000, São Paulo/SP, doravante denominada "
    "CONTRATANTE; e, de outro lado, BETA SERVIÇOS DE TECNOLOGIA S.A., pessoa jurídica "
    "de direito privado, inscrita no CNPJ sob o nº 98.765.432/0001-10, com sede na Rua "
    "das Flores, nº 250, Campinas/SP, doravante denominada CONTRATADA."
)
INTRO_2 = (
    "As partes acima identificadas têm, entre si, justo e acertado o presente Contrato "
    "de Prestação de Serviços, que se regerá pelas cláusulas seguintes e pelas "
    "condições descritas no presente instrumento."
)

OBJ_1 = (
    "O presente contrato tem como objeto a prestação, pela CONTRATADA, de serviços de "
    "consultoria técnica especializada em tecnologia da informação, compreendendo o "
    "diagnóstico, o planejamento e o acompanhamento da implantação de soluções de "
    "software na infraestrutura da CONTRATANTE."
)
OBJ_2 = (
    "Os serviços serão prestados nas dependências da CONTRATANTE ou remotamente, a "
    "critério das partes, observados os níveis de serviço definidos no Anexo I, que "
    "integra o presente instrumento para todos os fins de direito."
)
SCOPE = (
    "O escopo dos serviços compreende reuniões quinzenais de acompanhamento, "
    "elaboração de relatórios gerenciais mensais e suporte à tomada de decisão "
    "técnica, excluídas atividades de desenvolvimento de software sob demanda, que "
    "deverão ser objeto de aditivo específico."
)

OBRIG_A_1 = (
    "A CONTRATADA obriga-se a executar os serviços com zelo, diligência e observância "
    "das melhores práticas de mercado, alocando profissionais devidamente qualificados "
    "e mantendo a CONTRATANTE informada sobre o andamento das atividades."
)
OBRIG_A_2 = (
    "A CONTRATADA deverá observar todas as normas internas de segurança da informação "
    "da CONTRATANTE, bem como a legislação aplicável à proteção de dados pessoais, em "
    "especial a Lei nº 13.709/2018 (Lei Geral de Proteção de Dados Pessoais)."
)
OBRIG_A_3 = (
    "A CONTRATADA responsabiliza-se integralmente pelos encargos trabalhistas, "
    "previdenciários, fiscais e comerciais resultantes da execução do contrato, não se "
    "estabelecendo qualquer vínculo empregatício entre a CONTRATANTE e os "
    "profissionais alocados pela CONTRATADA."
)

OBRIG_B_1 = (
    "A CONTRATANTE obriga-se a fornecer à CONTRATADA todas as informações e acessos "
    "necessários à execução dos serviços, bem como a designar um responsável interno "
    "para o acompanhamento do contrato."
)
# Presente apenas na versão BASE — parágrafo removido na revisão.
REMOVED_PARA = (
    "A CONTRATANTE disponibilizará, sem ônus para a CONTRATADA, sala de reuniões "
    "equipada e acesso à rede corporativa de visitantes durante as atividades "
    "presenciais realizadas em suas dependências."
)
OBRIG_B_2 = (
    "A CONTRATANTE efetuará os pagamentos nos prazos e condições pactuados neste "
    "instrumento e comunicará à CONTRATADA, por escrito, qualquer irregularidade "
    "verificada na execução dos serviços."
)

PAY_INTRO = (
    "Pela prestação dos serviços objeto deste contrato, a CONTRATANTE pagará à "
    "CONTRATADA a remuneração mensal descrita na tabela de valores constante da "
    "cláusula 4.1, mediante apresentação de nota fiscal acompanhada do relatório de "
    "atividades do período."
)
PAY_PRAZO_BASE = (
    "O pagamento será realizado no prazo de 30 (trinta) dias corridos contados do "
    "recebimento da nota fiscal, mediante depósito em conta bancária de titularidade "
    "da CONTRATADA, valendo o comprovante de depósito como recibo de quitação."
)
PAY_PRAZO_REVISED = (
    "O pagamento será realizado no prazo de 45 (quarenta e cinco) dias corridos "
    "contados do recebimento da nota fiscal, mediante depósito em conta bancária de "
    "titularidade da CONTRATADA, valendo o comprovante de depósito como recibo de "
    "quitação."
)
TABLE_NOTE = (
    "Os valores constantes da tabela acima serão reajustados anualmente pela variação "
    "positiva do IPCA, ou por outro índice oficial que venha a substituí-lo, mediante "
    "simples apostilamento."
)

PEN_MULTA_BASE = (
    "O atraso injustificado no cumprimento das obrigações contratuais sujeitará a "
    "parte inadimplente à multa de 2% (dois por cento) sobre o valor mensal do "
    "contrato, sem prejuízo da apuração de perdas e danos."
)
PEN_MULTA_REVISED = (
    "O atraso injustificado no cumprimento das obrigações contratuais sujeitará a "
    "parte inadimplente à multa de 10% (dez por cento) sobre o valor mensal do "
    "contrato, sem prejuízo da apuração de perdas e danos."
)
PEN_2 = (
    "A aplicação da multa não impede a rescisão do contrato pela parte prejudicada, "
    "tampouco a cobrança de juros de mora de 1% (um por cento) ao mês e correção "
    "monetária sobre os valores em atraso."
)

CONF_1 = (
    "As partes obrigam-se a manter em sigilo todas as informações confidenciais a que "
    "tiverem acesso em razão do presente contrato, utilizando-as exclusivamente para a "
    "execução dos serviços contratados."
)
# Presente apenas na versão REVISADA — parágrafo novo de confidencialidade.
CONF_NEW = (
    "A obrigação de confidencialidade prevista nesta cláusula permanecerá vigente pelo "
    "período de 5 (cinco) anos contados do término ou rescisão do presente contrato, "
    "abrangendo também os colaboradores, prepostos e subcontratados de cada parte."
)

VIG_1 = (
    "O presente contrato vigorará pelo prazo de 12 (doze) meses contados de sua "
    "assinatura, podendo ser renovado por iguais e sucessivos períodos mediante termo "
    "aditivo assinado pelas partes."
)
# Parágrafo com APENAS formatação alterada na revisão (mesmo texto, negrito no meio).
VIG_FMT_PREFIX = "Qualquer das partes poderá rescindir o presente contrato, sem ônus, "
VIG_FMT_BOLD = "mediante notificação prévia e por escrito"
VIG_FMT_SUFFIX = " com antecedência mínima de 60 (sessenta) dias."

# Cláusula de foro — MOVIDA para o final do documento na revisão.
FORO_PARA = (
    "As partes elegem o foro da Comarca de São Paulo, Estado de São Paulo, com "
    "renúncia expressa a qualquer outro, por mais privilegiado que seja, para dirimir "
    "quaisquer controvérsias oriundas do presente contrato."
)

DISP_1 = (
    "O presente contrato representa o acordo integral entre as partes, substituindo "
    "todos os entendimentos anteriores, verbais ou escritos, e somente poderá ser "
    "alterado mediante termo aditivo assinado por ambas as partes."
)
DISP_2 = (
    "A tolerância de qualquer das partes quanto ao descumprimento de obrigação da "
    "outra não constituirá novação, renúncia ou alteração do pactuado, podendo a parte "
    "tolerante exigir o cumprimento da obrigação a qualquer tempo."
)
CLOSING = (
    "E por estarem assim justas e contratadas, as partes assinam o presente "
    "instrumento em 2 (duas) vias de igual teor e forma, na presença das testemunhas "
    "abaixo qualificadas."
)


def _contract_table_rows(revised: bool) -> List[List[str]]:
    rows = [
        ["Serviço", "Periodicidade", "Valor"],
        [
            "Consultoria técnica",
            "Mensal",
            "R$ 7.500,00" if revised else "R$ 5.000,00",  # célula alterada
        ],
        ["Suporte remoto", "Mensal", "R$ 1.200,00"],
        ["Treinamento da equipe", "Trimestral", "R$ 3.800,00"],
    ]
    if revised:
        # Linha nova na tabela.
        rows.append(["Auditoria de segurança", "Semestral", "R$ 4.500,00"])
    return rows


def _build_contract_docx(path: str, revised: bool) -> None:
    doc = DocxDocument()
    doc.add_heading(CONTRACT_TITLE, level=0)
    doc.add_paragraph(VERSION_LINE_REVISED if revised else VERSION_LINE_BASE)
    doc.add_paragraph(INTRO_1)
    doc.add_paragraph(INTRO_2)

    doc.add_heading("1. Do Objeto", level=1)
    doc.add_paragraph(OBJ_1)
    doc.add_paragraph(OBJ_2)
    doc.add_heading("1.1. Do Escopo dos Serviços", level=2)
    doc.add_paragraph(SCOPE)

    doc.add_heading("2. Das Obrigações da Contratada", level=1)
    doc.add_paragraph(OBRIG_A_1)
    doc.add_paragraph(OBRIG_A_2)
    doc.add_paragraph(OBRIG_A_3)

    doc.add_heading("3. Das Obrigações da Contratante", level=1)
    doc.add_paragraph(OBRIG_B_1)
    if not revised:
        # Parágrafo removido na versão revisada.
        doc.add_paragraph(REMOVED_PARA)
    doc.add_paragraph(OBRIG_B_2)

    doc.add_heading("4. Do Pagamento", level=1)
    doc.add_paragraph(PAY_INTRO)
    doc.add_paragraph(PAY_PRAZO_REVISED if revised else PAY_PRAZO_BASE)
    doc.add_heading("4.1. Da Tabela de Valores", level=2)
    rows = _contract_table_rows(revised)
    table = doc.add_table(rows=len(rows), cols=3)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass  # estilo ausente no template não impede a geração
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = value
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    doc.add_paragraph(TABLE_NOTE)

    doc.add_heading("5. Das Penalidades", level=1)
    doc.add_paragraph(PEN_MULTA_REVISED if revised else PEN_MULTA_BASE)
    doc.add_paragraph(PEN_2)

    doc.add_heading("6. Da Confidencialidade", level=1)
    doc.add_paragraph(CONF_1)
    if revised:
        # Parágrafo novo de confidencialidade.
        doc.add_paragraph(CONF_NEW)

    doc.add_heading("7. Da Vigência e da Rescisão", level=1)
    doc.add_paragraph(VIG_1)
    if revised:
        # Mesmo texto, negrito adicionado no trecho central (só formatação).
        p = doc.add_paragraph()
        p.add_run(VIG_FMT_PREFIX)
        bold_run = p.add_run(VIG_FMT_BOLD)
        bold_run.bold = True
        p.add_run(VIG_FMT_SUFFIX)
    else:
        doc.add_paragraph(VIG_FMT_PREFIX + VIG_FMT_BOLD + VIG_FMT_SUFFIX)

    doc.add_heading("8. Das Disposições Gerais", level=1)
    if not revised:
        doc.add_paragraph(FORO_PARA)
    doc.add_paragraph(DISP_1)
    doc.add_paragraph(DISP_2)
    doc.add_paragraph(CLOSING)
    if revised:
        # Cláusula de foro movida para o final do documento.
        doc.add_paragraph(FORO_PARA)

    doc.save(path)


# ---------------------------------------------------------------------------
# Política de Privacidade (PDF via fitz)
# ---------------------------------------------------------------------------

POLICY_TITLE = "Política de Privacidade"
POLICY_DATE_BASE = "Última atualização: 10 de janeiro de 2026."
POLICY_DATE_REVISED = "Última atualização: 05 de julho de 2026."

POLICY_USO_BASE = (
    "Os dados coletados são utilizados para viabilizar a prestação dos serviços "
    "contratados, processar pagamentos, enviar comunicações operacionais e aprimorar "
    "a experiência do usuário na plataforma."
)
POLICY_USO_REVISED = (
    "Os dados coletados são utilizados para viabilizar a prestação dos serviços "
    "contratados, processar pagamentos, enviar comunicações operacionais, realizar "
    "análises estatísticas agregadas e aprimorar continuamente a experiência do "
    "usuário na plataforma."
)
POLICY_SEG_NEW = (
    "Realizamos auditorias periódicas de segurança e testes de intrusão conduzidos "
    "por equipes independentes, e mantemos um plano de resposta a incidentes com "
    "prazos definidos de notificação aos titulares e à Autoridade Nacional de "
    "Proteção de Dados."
)


def _policy_items(revised: bool) -> List[Tuple[str, str]]:
    items: List[Tuple[str, str]] = [
        ("h1", POLICY_TITLE),
        ("body", POLICY_DATE_REVISED if revised else POLICY_DATE_BASE),
        ("h2", "1. Introdução"),
        (
            "body",
            "Esta Política de Privacidade descreve como a Compare Docs Tecnologia "
            "Ltda. coleta, utiliza, armazena e protege os dados pessoais dos usuários "
            "de seus produtos e serviços, em conformidade com a Lei Geral de Proteção "
            "de Dados Pessoais (Lei nº 13.709/2018).",
        ),
        (
            "body",
            "Ao utilizar nossos serviços, o usuário declara ter lido e compreendido "
            "os termos desta política. Recomendamos a leitura atenta e integral deste "
            "documento antes de qualquer utilização da plataforma.",
        ),
        ("h2", "2. Dados Coletados"),
        (
            "body",
            "Coletamos dados fornecidos diretamente pelo usuário no momento do "
            "cadastro, tais como nome completo, endereço de e-mail, telefone de "
            "contato e dados de faturamento, quando aplicável ao plano contratado.",
        ),
        (
            "body",
            "Também coletamos automaticamente informações técnicas relacionadas ao "
            "uso da plataforma, incluindo endereço IP, tipo de navegador, sistema "
            "operacional, páginas acessadas e registros de data e hora de acesso.",
        ),
        (
            "body",
            "Não coletamos dados pessoais sensíveis, tais como origem racial ou "
            "étnica, convicção religiosa, opinião política ou dados referentes à "
            "saúde, salvo quando estritamente necessário e mediante consentimento "
            "específico e destacado do titular.",
        ),
        ("h2", "3. Uso das Informações"),
        ("body", POLICY_USO_REVISED if revised else POLICY_USO_BASE),
        (
            "body",
            "Comunicações de marketing somente serão enviadas mediante consentimento "
            "prévio do usuário, que poderá revogá-lo a qualquer momento por meio do "
            "link de descadastramento presente em cada mensagem enviada.",
        ),
        ("h2", "4. Compartilhamento de Dados"),
        (
            "body",
            "Não vendemos, alugamos ou comercializamos dados pessoais de usuários. O "
            "compartilhamento de dados com terceiros ocorre apenas nas hipóteses "
            "descritas nesta política ou mediante obrigação legal ou regulatória.",
        ),
        (
            "body",
            "Podemos compartilhar dados com prestadores de serviços essenciais à "
            "operação, tais como provedores de infraestrutura em nuvem e "
            "processadores de pagamento, sempre mediante contratos que assegurem "
            "padrões adequados de proteção e confidencialidade.",
        ),
        ("h2", "5. Segurança da Informação"),
        (
            "body",
            "Adotamos medidas técnicas e organizacionais apropriadas para proteger os "
            "dados pessoais contra acessos não autorizados, perda, alteração ou "
            "destruição, incluindo criptografia em trânsito e em repouso.",
        ),
    ]
    if revised:
        # Parágrafo inserido na revisão.
        items.append(("body", POLICY_SEG_NEW))
    items.extend(
        [
            (
                "body",
                "O acesso aos dados pessoais é restrito aos colaboradores que deles "
                "necessitam para o exercício de suas funções, os quais se sujeitam a "
                "obrigações contratuais de confidencialidade e sigilo profissional.",
            ),
            ("h2", "6. Direitos do Titular"),
            (
                "body",
                "Nos termos da legislação aplicável, o titular dos dados poderá "
                "solicitar a confirmação da existência de tratamento, o acesso aos "
                "dados, a correção de dados incompletos ou desatualizados, a "
                "anonimização, o bloqueio ou a eliminação de dados desnecessários ou "
                "excessivos.",
            ),
            (
                "body",
                "As solicitações serão respondidas no prazo de 15 (quinze) dias "
                "contados do recebimento, prazo que poderá ser prorrogado mediante "
                "justificativa expressa comunicada ao titular.",
            ),
            ("h2", "7. Contato"),
            (
                "body",
                "Em caso de dúvidas sobre esta política ou sobre o tratamento de "
                "dados pessoais, o usuário poderá contatar o Encarregado de Proteção "
                "de Dados pelo e-mail privacidade@comparedocs.com.br.",
            ),
        ]
    )
    return items


_PDF_STYLES = {
    # kind: (fontname, fontsize, espaço após o bloco)
    "h1": ("hebo", 18.0, 14.0),
    "h2": ("hebo", 14.0, 10.0),
    "body": ("helv", 11.0, 9.0),
}


def _build_policy_pdf(path: str, revised: bool) -> None:
    margin = 56.0
    page_rect = fitz.paper_rect("a4")
    doc = fitz.open()
    page = doc.new_page(width=page_rect.width, height=page_rect.height)
    y = margin
    for kind, text in _policy_items(revised):
        fontname, fontsize, gap = _PDF_STYLES[kind]
        while True:
            rect = fitz.Rect(margin, y, page_rect.width - margin, page_rect.height - margin)
            rc = -1.0
            if rect.height >= fontsize * 1.6:
                rc = page.insert_textbox(
                    rect, text, fontname=fontname, fontsize=fontsize, align=0
                )
            if rc >= 0:
                y = y + (rect.height - rc) + gap
                break
            if y <= margin + 0.5:
                doc.close()
                raise ValueError(
                    "Não foi possível posicionar o texto no PDF de amostra: %r"
                    % text[:60]
                )
            page = doc.new_page(width=page_rect.width, height=page_rect.height)
            y = margin
    doc.save(path)
    doc.close()


# ---------------------------------------------------------------------------
# Proposta Comercial (par com nomes diferentes — pareamento fuzzy)
# ---------------------------------------------------------------------------

PROPOSAL_INTRO = (
    "Apresentamos a proposta comercial para o projeto de implantação do sistema de "
    "gestão documental, conforme escopo discutido em reunião realizada em 15 de junho "
    "de 2026 com a equipe técnica da contratante."
)
PROPOSAL_VALUE_BASE = (
    "O investimento total para a execução do projeto é de R$ 25.000,00 (vinte e cinco "
    "mil reais), incluindo licenciamento, implantação e treinamento da equipe."
)
PROPOSAL_VALUE_REVISED = (
    "O investimento total para a execução do projeto é de R$ 28.500,00 (vinte e oito "
    "mil e quinhentos reais), incluindo licenciamento, implantação e treinamento da "
    "equipe."
)
PROPOSAL_DEADLINE_BASE = (
    "O prazo de entrega estimado é de 60 (sessenta) dias corridos contados da "
    "assinatura do contrato, condicionado à disponibilidade dos ambientes de "
    "homologação da contratante."
)
PROPOSAL_DEADLINE_REVISED = (
    "O prazo de entrega estimado é de 90 (noventa) dias corridos contados da "
    "assinatura do contrato, condicionado à disponibilidade dos ambientes de "
    "homologação da contratante."
)
PROPOSAL_VALIDITY = (
    "Esta proposta é válida por 30 (trinta) dias contados da data de seu envio, e os "
    "valores apresentados não incluem tributos incidentes sobre o faturamento."
)


def _build_proposal_docx(path: str, revised: bool) -> None:
    doc = DocxDocument()
    doc.add_heading("Proposta Comercial", level=0)
    doc.add_heading("1. Apresentação", level=1)
    doc.add_paragraph(PROPOSAL_INTRO)
    doc.add_heading("2. Investimento", level=1)
    doc.add_paragraph(PROPOSAL_VALUE_REVISED if revised else PROPOSAL_VALUE_BASE)
    doc.add_heading("3. Prazo de Entrega", level=1)
    doc.add_paragraph(PROPOSAL_DEADLINE_REVISED if revised else PROPOSAL_DEADLINE_BASE)
    doc.add_heading("4. Validade", level=1)
    doc.add_paragraph(PROPOSAL_VALIDITY)
    doc.save(path)


# ---------------------------------------------------------------------------
# Orçamento (Excel — comparação de planilhas)
# ---------------------------------------------------------------------------


def _build_budget_xlsx(path: str, revised: bool) -> None:
    wb = Workbook()
    costs = wb.active
    costs.title = "Custos"
    costs.append(["Item", "Qtd", "Valor Unit.", "Total"])
    costs.append(["Consultoria", 10, 500, 5000])
    costs.append(["Licença anual", 1, 2500, 2500])
    if revised:
        costs.append(["Suporte mensal", 12, 200, 2400])
        costs["C3"] = 600
    else:
        costs.append(["Treinamento", 2, 800, 1600])

    summary = wb.create_sheet("Resumo")
    summary["A1"] = "Total do projeto"
    summary["B1"] = 10900 if revised else 9100
    summary["A2"] = "Prazo (dias)"
    summary["B2"] = 90 if revised else 60
    wb.save(path)


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main() -> None:
    """Apaga e recria tests/samples/ com os três pares de amostra."""
    shutil.rmtree(SAMPLES_DIR, ignore_errors=True)
    os.makedirs(BASE_DIR)
    os.makedirs(REVISED_DIR)

    created: List[str] = []

    contract_base = os.path.join(BASE_DIR, CONTRACT_NAME)
    contract_revised = os.path.join(REVISED_DIR, CONTRACT_NAME)
    _build_contract_docx(contract_base, revised=False)
    _build_contract_docx(contract_revised, revised=True)
    created.extend([contract_base, contract_revised])

    policy_base = os.path.join(BASE_DIR, POLICY_NAME)
    policy_revised = os.path.join(REVISED_DIR, POLICY_NAME)
    _build_policy_pdf(policy_base, revised=False)
    _build_policy_pdf(policy_revised, revised=True)
    created.extend([policy_base, policy_revised])

    proposal_base = os.path.join(BASE_DIR, PROPOSAL_BASE_NAME)
    proposal_revised = os.path.join(REVISED_DIR, PROPOSAL_REVISED_NAME)
    _build_proposal_docx(proposal_base, revised=False)
    _build_proposal_docx(proposal_revised, revised=True)
    created.extend([proposal_base, proposal_revised])

    budget_base = os.path.join(BASE_DIR, BUDGET_NAME)
    budget_revised = os.path.join(REVISED_DIR, BUDGET_NAME)
    _build_budget_xlsx(budget_base, revised=False)
    _build_budget_xlsx(budget_revised, revised=True)
    created.extend([budget_base, budget_revised])

    for path in created:
        if not os.path.isfile(path) or os.path.getsize(path) == 0:
            raise ValueError(
                "Falha ao gerar a amostra %r: arquivo ausente ou vazio." % path
            )

    print("Amostras geradas em %s:" % SAMPLES_DIR)
    for path in created:
        print("  - %s (%d bytes)" % (os.path.relpath(path, SAMPLES_DIR), os.path.getsize(path)))


if __name__ == "__main__":
    main()
