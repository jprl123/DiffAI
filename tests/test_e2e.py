"""Testes ponta-a-ponta do Compare-Docs.

Gera amostras conhecidas, executa o pipeline completo (extração → comparação
→ saídas) e valida contagens mínimas e tipos de mudança esperados.
"""
from __future__ import annotations

import os
import shutil
import sys
import tempfile
import unittest

from app.batch import pair_files
from app.engine.compare import compare_documents
from app.extract.loader import load_document
from app.jobs import JobManager
from app.models import BlockKind, Category, ChangeType
from app.output.redline_docx import write_redline_docx
from app.output.redline_docx_inplace import write_redline_docx_inplace
from app.output.redline_pdf import write_redline_pdf
from app.output.report import write_html_report, write_json_report, write_xlsx_report
from tests.make_samples import (
    BASE_DIR,
    CONTRACT_NAME,
    POLICY_NAME,
    PROPOSAL_BASE_NAME,
    REVISED_DIR,
    SAMPLES_DIR,
    BUDGET_NAME,
    main as make_samples,
)

TESTS_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(TESTS_DIR)

RIVIO_FIXTURE_DIR = os.path.join(TESTS_DIR, "fixtures", "rivio")
RIVIO_NAME = "Rivio – Series Seed Share Preference Share Purchase Agreement.docx"
RIVIO_BASE_DIR = os.path.join(RIVIO_FIXTURE_DIR, "base")
RIVIO_REVISED_DIR = os.path.join(RIVIO_FIXTURE_DIR, "revised")


def _isolated_manager() -> JobManager:
    """JobManager com histórico descartável — testes não podem poluir o
    histórico real do usuário (~/.comparedocs/history.json)."""
    from app.history import HistoryStore

    tmp = tempfile.mkdtemp(prefix="comparedocs-test-history-")
    return JobManager(history_store=HistoryStore(path=os.path.join(tmp, "h.json")))


class TestEndToEnd(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        make_samples()

    def _compare_pair(self, base_name: str, revised_name: str):
        base_path = os.path.join(BASE_DIR, base_name)
        revised_path = os.path.join(REVISED_DIR, revised_name)
        base_doc = load_document(base_path)
        revised_doc = load_document(revised_path)
        return compare_documents(base_doc, revised_doc), base_path, revised_path

    def test_samples_exist(self) -> None:
        self.assertTrue(os.path.isdir(SAMPLES_DIR))
        for name in (CONTRACT_NAME, POLICY_NAME, PROPOSAL_BASE_NAME):
            self.assertTrue(os.path.isfile(os.path.join(BASE_DIR, name)))

    def test_contract_docx_pipeline(self) -> None:
        result, base_path, revised_path = self._compare_pair(CONTRACT_NAME, CONTRACT_NAME)
        self.assertGreaterEqual(result.stats.total_changes, 7)
        self.assertGreaterEqual(result.stats.content_changes, 4)
        self.assertGreaterEqual(result.stats.formatting_changes, 1)
        self.assertGreaterEqual(result.stats.noise_changes, 1)
        self.assertGreaterEqual(result.stats.table_changes, 1)
        self.assertGreaterEqual(result.stats.moves, 1)

        categories = {c.category for c in result.changes}
        self.assertIn(Category.CONTENT, categories)
        self.assertIn(Category.FORMATTING, categories)
        self.assertTrue(
            Category.NOISE_VERSION in categories or Category.NOISE_DATE in categories
        )

        types = {c.change_type for c in result.changes}
        self.assertIn(ChangeType.INSERT, types)
        self.assertIn(ChangeType.DELETE, types)
        self.assertIn(ChangeType.MOVE, types)

        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-")
        try:
            pdf_path = os.path.join(out_dir, "redline.pdf")
            write_redline_pdf(result, pdf_path)
            self.assertGreater(os.path.getsize(pdf_path), 1000)

            write_redline_docx(result, os.path.join(out_dir, "redline.docx"))
            write_html_report(result, os.path.join(out_dir, "report.html"))
            write_xlsx_report(result, os.path.join(out_dir, "report.xlsx"))
            write_json_report(result, os.path.join(out_dir, "report.json"))

            self.assertTrue(os.path.isfile(os.path.join(out_dir, "redline.docx")))
            self.assertTrue(os.path.isfile(os.path.join(out_dir, "report.html")))
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

        self.assertTrue(result.base_path.endswith(CONTRACT_NAME))
        self.assertTrue(result.compare_path.endswith(CONTRACT_NAME))

    def test_policy_pdf_pipeline(self) -> None:
        result, _, _ = self._compare_pair(POLICY_NAME, POLICY_NAME)
        self.assertGreaterEqual(result.stats.total_changes, 2)
        self.assertGreaterEqual(result.stats.content_changes, 1)
        self.assertTrue(len(result.render_blocks) > 0)

        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-pdf-")
        try:
            pdf_path = os.path.join(out_dir, "policy.pdf")
            write_redline_pdf(result, pdf_path)
            self.assertGreater(os.path.getsize(pdf_path), 500)
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def test_batch_pairing_fuzzy_names(self) -> None:
        pairs, unmatched_base, unmatched_compare = pair_files(BASE_DIR, REVISED_DIR)
        self.assertEqual(len(pairs), 4)
        self.assertEqual(unmatched_base, [])
        self.assertEqual(unmatched_compare, [])

        proposal_pairs = [
            p for p in pairs if PROPOSAL_BASE_NAME in os.path.basename(p[0])
        ]
        self.assertEqual(len(proposal_pairs), 1)

        result, _, _ = self._compare_pair(PROPOSAL_BASE_NAME, "Proposta Comercial v2 final.docx")
        self.assertGreaterEqual(result.stats.total_changes, 2)
        self.assertGreaterEqual(result.stats.content_changes, 2)

    def test_budget_xlsx_pipeline(self) -> None:
        result, base_path, revised_path = self._compare_pair(BUDGET_NAME, BUDGET_NAME)
        self.assertGreaterEqual(result.stats.total_changes, 2)
        self.assertGreaterEqual(result.stats.table_changes, 1)

        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-xlsx-")
        try:
            from app.output.redline_xlsx import write_redline_xlsx

            redline_path = os.path.join(out_dir, "redline.xlsx")
            write_redline_xlsx(base_path, revised_path, redline_path)
            self.assertGreater(os.path.getsize(redline_path), 1000)
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def test_xlsx_job_manager(self) -> None:
        manager = _isolated_manager()
        base_path = os.path.join(BASE_DIR, BUDGET_NAME)
        revised_path = os.path.join(REVISED_DIR, BUDGET_NAME)
        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-xlsx-job-")
        try:
            job_id = manager.create_job(
                [(base_path, revised_path)],
                {
                    "changed_pages_only": False,
                    "export_docx": False,
                    "reports": ["html"],
                    "output_dir": out_dir,
                },
            )
            import time

            deadline = time.time() + 30.0
            job = None
            while time.time() < deadline:
                job = manager.get_job(job_id)
                if job and job.get("status") in ("done", "error"):
                    break
                time.sleep(0.2)

            self.assertIsNotNone(job)
            self.assertEqual(job["status"], "done")
            item = job["items"][0]
            self.assertEqual(item["status"], "ok")
            self.assertIn("redline_xlsx", item["outputs"])
            self.assertTrue(os.path.isfile(item["outputs"]["redline_xlsx"]))
            self.assertNotIn("pdf", item["outputs"])
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def test_rivio_real_world_docx(self) -> None:
        """Documento real (Series Seed SPA) — não é apagado por make_samples."""
        base_path = os.path.join(RIVIO_BASE_DIR, RIVIO_NAME)
        revised_path = os.path.join(RIVIO_REVISED_DIR, RIVIO_NAME)
        self.assertTrue(os.path.isfile(base_path), "Fixture Rivio base ausente")
        self.assertTrue(os.path.isfile(revised_path), "Fixture Rivio revisado ausente")

        result, _, _ = self._compare_pair_from_paths(base_path, revised_path)
        self.assertGreaterEqual(result.stats.total_changes, 10)
        self.assertGreaterEqual(result.stats.content_changes, 10)
        # Ground truth verificado: o par Rivio tem 172 parágrafos idênticos
        # únicos e ZERO inversões de ordem — qualquer "movimentação" aqui é
        # falso positivo do alinhador.
        self.assertEqual(result.stats.moves, 0)
        self.assertGreater(len(result.render_blocks), 50)

        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-rivio-")
        try:
            docx_path = os.path.join(out_dir, "redline.docx")
            write_redline_docx_inplace(result, revised_path, docx_path)
            self.assertGreater(os.path.getsize(docx_path), 50_000)

            pdf_path = os.path.join(out_dir, "redline.pdf")
            write_redline_pdf(result, pdf_path)
            self.assertGreater(os.path.getsize(pdf_path), 1000)
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def _compare_pair_from_paths(self, base_path: str, revised_path: str):
        base_doc = load_document(base_path)
        revised_doc = load_document(revised_path)
        return compare_documents(base_doc, revised_doc), base_path, revised_path

    def test_inplace_docx_summary_last_page(self) -> None:
        """A síntese ('Summary of Changes' + marca do app) fecha o DOCX fiel,
        em seção própria (sem herdar cabeçalho/rodapé do documento)."""
        from docx import Document as DocxDocument

        result, _, revised_path = self._compare_pair(CONTRACT_NAME, CONTRACT_NAME)
        result.compared_at = "2026-07-08T12:00:00"
        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-summary-")
        try:
            docx_path = os.path.join(out_dir, "redline.docx")
            write_redline_docx_inplace(result, revised_path, docx_path)
            doc = DocxDocument(docx_path)
            texts = [p.text for p in doc.paragraphs]
            self.assertIn("Summary of Changes", texts)
            self.assertIn("COMPARE DOCS", texts)
            # A síntese vem DEPOIS do último parágrafo de conteúdo.
            summary_idx = texts.index("Summary of Changes")
            self.assertGreater(summary_idx, len(texts) // 2)
            # Seção própria, com rodapé desvinculado do documento.
            self.assertGreaterEqual(len(doc.sections), 2)
            self.assertFalse(doc.sections[-1].footer.is_linked_to_previous)
            # A tabela de métricas existe e cobre os totais.
            labels = [
                row.cells[0].text for table in doc.tables for row in table.rows
            ]
            self.assertIn("Total de alterações", labels)
            self.assertIn("Movimentações", labels)
            # Síntese enxuta (decisão do usuário 2026-07-12): sem rotineiras,
            # formatação, tabelas ou imagens.
            self.assertNotIn("Mudanças rotineiras (ruído)", labels)
            self.assertNotIn("Formatação", labels)
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def test_history_persistence(self) -> None:
        """Cada par processado vira entrada persistente; /api/open confia nela."""
        import time

        from app.history import HistoryStore

        tmp = tempfile.mkdtemp(prefix="comparedocs-e2e-history-")
        try:
            store = HistoryStore(path=os.path.join(tmp, "history.json"))
            manager = JobManager(history_store=store)
            base_path = os.path.join(BASE_DIR, PROPOSAL_BASE_NAME)
            revised_path = os.path.join(REVISED_DIR, "Proposta Comercial v2 final.docx")
            out_dir = os.path.join(tmp, "out")
            job_id = manager.create_job(
                [(base_path, revised_path)], {"output_dir": out_dir}
            )
            deadline = time.time() + 120
            while manager.get_job(job_id)["status"] == "running":
                self.assertLess(time.time(), deadline, "job não terminou a tempo")
                time.sleep(0.2)

            entries = store.list_entries()
            self.assertEqual(len(entries), 1)
            entry = entries[0]
            self.assertEqual(entry["status"], "ok")
            self.assertEqual(entry["base_name"], PROPOSAL_BASE_NAME)
            self.assertIn("pdf", entry["outputs"])
            self.assertGreaterEqual((entry["stats"] or {})["total_changes"], 1)

            # Outra instância (novo processo/sessão) lê o mesmo arquivo e
            # reconhece os caminhos gerados — base do /api/open pós-restart.
            fresh = HistoryStore(path=store.path)
            self.assertTrue(fresh.path_known(entry["outputs"]["pdf"]))
            self.assertFalse(fresh.path_known("/etc/passwd"))

            from app.history_result import result_dict_for_entry

            result_data = result_dict_for_entry(entry)
            self.assertIsNotNone(result_data)
            self.assertGreater(len(result_data.get("changes") or []), 0)

            self.assertTrue(fresh.remove_entry(entry["id"]))
            self.assertEqual(fresh.list_entries(), [])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    def test_escorrega_real_world_matches_word_compare(self) -> None:
        """Contrato real validado contra o redline do Word (regressões do
        motor encontradas em 2026-07-17): sem texto idêntico tachado e
        reinserido, sem falso-movido, e diff fino em pontuação."""
        import re

        base_path = os.path.join(
            TESTS_DIR, "fixtures", "escorrega", "base",
            "Opção_Compra_Imóvel_ESCORREGA_Piauí_23.02.2026 (1).docx",
        )
        revised_path = os.path.join(
            TESTS_DIR, "fixtures", "escorrega", "revised",
            "10b & Passuarê I Opção de Compra Imóvel - Fazenda Escorrega (1).docx",
        )
        if not os.path.isfile(base_path) or not os.path.isfile(revised_path):
            self.skipTest("fixture escorrega ausente")

        result, _, _ = self._compare_pair_from_paths(base_path, revised_path)

        # Ground truth: zero inversões de ordem entre os dois documentos.
        self.assertEqual(result.stats.moves, 0)

        espolio_found = False
        for rb in result.render_blocks:
            dels = "".join(f.text for f in rb.fragments if f.op == "delete")
            ins = "".join(f.text for f in rb.fragments if f.op == "insert")
            # Nenhum bloco pode tachar e reinserir texto idêntico (fantasma).
            if dels.strip() and ins.strip():
                self.assertNotEqual(
                    re.sub(r"\s+", " ", dels).strip(),
                    re.sub(r"\s+", " ", ins).strip(),
                    "diff fantasma: texto idêntico excluído e reinserido",
                )
            text = "".join(f.text for f in rb.fragments)
            if "ESPÓLIO DE MANOEL" in text:
                espolio_found = True
                # O nome inteiro deve permanecer equal (o Word marca só o ponto).
                equal_text = "".join(
                    f.text for f in rb.fragments if f.op == "equal"
                )
                self.assertIn("ESPÓLIO DE MANOEL JOSÉ DE SANTANA", equal_text)
        self.assertTrue(espolio_found, "parágrafo do ESPÓLIO não encontrado")

    def test_memorandum_definitions_table(self) -> None:
        """Memorandum real (tabela de definições, 40→36 linhas) — regressões
        reportadas pelo usuário em 2026-07-12 contra o Word Compare:
        exclusão de linha DEVE aparecer; termos idênticos não podem ser
        tachados+reinseridos; definição com 1 palavra alterada é modify."""
        base_path = os.path.join(
            TESTS_DIR, "fixtures", "memorandum", "base",
            "[Original] XXX - Series Seed - Memorandum and Articles [VF].docx",
        )
        revised_path = os.path.join(
            TESTS_DIR, "fixtures", "memorandum", "revised",
            "[New Version] XXX - Series Seed - Memorandum and Articles [VF].docx",
        )
        if not os.path.isfile(base_path) or not os.path.isfile(revised_path):
            self.skipTest("fixture memorandum ausente")

        result, _, _ = self._compare_pair_from_paths(base_path, revised_path)
        table = next(
            rb for rb in result.render_blocks if rb.kind == BlockKind.TABLE
        )

        def row_term(row):
            return "".join(f.text for f in row[0]) if row else ""

        def cell_marked_text(row, op):
            return " ".join(
                f.text for cell in row for f in cell if f.op == op
            )

        rows = list(zip(table.rows, table.row_ops))

        # 1. Exclusão de linha inteira aparece (caso "Seal"/"CEO").
        deleted_terms = [row_term(r) for r, op in rows if op == "delete"]
        self.assertTrue(
            any("Seal" in t for t in deleted_terms),
            "linha excluída 'Seal' não marcada como delete: %r" % deleted_terms,
        )

        # 2. Palavra idêntica nunca tachada+reinserida (casos including/Company).
        for row, op in rows:
            if op != "modify":
                continue
            term = row_term(row)
            if "Affiliate" in term:
                self.assertNotIn("including", cell_marked_text(row, "insert"))
                self.assertNotIn("including", cell_marked_text(row, "delete").replace(", without limitation,", ""))
            if "Articles" in term and "”" in term:
                self.assertNotIn("Company", cell_marked_text(row, "delete"))
                self.assertNotIn("Company", cell_marked_text(row, "insert"))

        # 3. Definições quase iguais são MODIFY, não delete+insert.
        preseed_ops = {
            row_term(r)[:40]: op for r, op in rows if "Pre-Seed-3" in row_term(r) or "Pre-Seed-4" in row_term(r)
        }
        self.assertTrue(preseed_ops, "linhas Pre-Seed não encontradas")
        self.assertTrue(
            all(op == "modify" for op in preseed_ops.values()),
            "Pre-Seed-3/4 deveriam ser modify: %r" % preseed_ops,
        )

        # 4. Renumeração automática de cláusulas ((a)→(b) na 6.1) é detectada
        # e nunca classificada como ruído.
        renum = [c for c in result.changes if "Numeração alterada" in c.summary]
        self.assertTrue(
            any("(a) para (b)" in c.summary for c in renum),
            "renumeração (a)→(b) da cláusula 6.1 não detectada",
        )
        for c in renum:
            self.assertEqual(c.category, Category.CONTENT)

        # 5. Regra dos 30%: palavra substancialmente trocada sai inteira —
        # nenhum fragmento pode misturar pedaços de "result of"/"resulting".
        for rb in result.render_blocks:
            for f in rb.fragments:
                if f.op == "insert":
                    self.assertNotIn("ofing", f.text, "letras recicladas no diff")

    def test_exec_summary_one_page(self) -> None:
        """Resumo executivo: sempre 1 página, com síntese e destaques."""
        import fitz

        from app.output.exec_summary import write_exec_summary_pdf

        result, _, _ = self._compare_pair(CONTRACT_NAME, CONTRACT_NAME)
        result.compared_at = "2026-07-17T10:00:00"
        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-execsum-")
        try:
            out = os.path.join(out_dir, "resumo.pdf")
            write_exec_summary_pdf(result, out)
            doc = fitz.open(out)
            self.assertEqual(doc.page_count, 1)
            text = doc[0].get_text()
            self.assertIn("Resumo Executivo", text)
            self.assertIn("Síntese", text)
            self.assertTrue(
                "Comparação entre" in text or "mudança" in text.lower(),
                "síntese executiva ausente",
            )
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)

    def test_batch_preview_detailed(self) -> None:
        """Prévia do lote informa o método de pareamento de cada par."""
        from app.batch import pair_files_detailed

        pairs, unmatched_base, unmatched_compare = pair_files_detailed(
            BASE_DIR, REVISED_DIR
        )
        self.assertGreaterEqual(len(pairs), 3)
        methods = {p["method"] for p in pairs}
        self.assertTrue(methods <= {"nome", "similaridade", "conteúdo"})
        proposal = [p for p in pairs if p["base_name"] == PROPOSAL_BASE_NAME]
        self.assertEqual(len(proposal), 1)
        self.assertEqual(proposal[0]["method"], "similaridade")

    def test_scanned_pdf_rejected(self) -> None:
        """PDF digitalizado (sem camada de texto) falha com mensagem clara."""
        import fitz

        tmp = tempfile.mkdtemp(prefix="comparedocs-e2e-scan-")
        try:
            path = os.path.join(tmp, "escaneado.pdf")
            doc = fitz.open()
            page = doc.new_page()
            pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 100, 100), False)
            pix.clear_with(180)
            page.insert_image(fitz.Rect(50, 50, 550, 750), pixmap=pix)
            doc.save(path)
            doc.close()
            with self.assertRaises(ValueError) as ctx:
                load_document(path)
            self.assertIn("digitalizado", str(ctx.exception))
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    def test_batch_pairing_by_content(self) -> None:
        """Arquivos com nomes SEM relação pareiam pelo conteúdo do texto."""
        from docx import Document as DocxDocument

        tmp = tempfile.mkdtemp(prefix="comparedocs-e2e-content-")
        try:
            base_dir = os.path.join(tmp, "base")
            rev_dir = os.path.join(tmp, "rev")
            os.makedirs(base_dir)
            os.makedirs(rev_dir)

            def make(path: str, extra: str = "") -> None:
                doc = DocxDocument()
                for i in range(8):
                    doc.add_paragraph(
                        "Cláusula %d. As partes acordam em cooperar no projeto "
                        "Alfa, respeitando prazos e valores deste instrumento." % (i + 1)
                    )
                if extra:
                    doc.add_paragraph(extra)
                doc.save(path)

            make(os.path.join(base_dir, "documento_final_assinado.docx"))
            make(
                os.path.join(rev_dir, "ACT-2026-jur-07 (revisado).docx"),
                extra="Cláusula nova de confidencialidade.",
            )
            # Documento de conteúdo distinto NÃO deve parear.
            other = DocxDocument()
            other.add_paragraph("Relatório de despesas de viagem e reembolsos.")
            other.save(os.path.join(base_dir, "despesas.docx"))

            pairs, unmatched_base, unmatched_compare = pair_files(base_dir, rev_dir)
            self.assertEqual(len(pairs), 1)
            self.assertEqual(
                os.path.basename(pairs[0][0]), "documento_final_assinado.docx"
            )
            self.assertEqual(
                os.path.basename(pairs[0][1]), "ACT-2026-jur-07 (revisado).docx"
            )
            self.assertEqual(unmatched_base, ["despesas.docx"])
            self.assertEqual(unmatched_compare, [])
        finally:
            shutil.rmtree(tmp, ignore_errors=True)

    def test_job_manager_integration(self) -> None:
        manager = _isolated_manager()
        base_path = os.path.join(BASE_DIR, PROPOSAL_BASE_NAME)
        revised_path = os.path.join(REVISED_DIR, "Proposta Comercial v2 final.docx")
        out_dir = tempfile.mkdtemp(prefix="comparedocs-e2e-job-")
        try:
            job_id = manager.create_job(
                [(base_path, revised_path)],
                {
                    "changed_pages_only": False,
                    "export_docx": True,
                    "reports": ["html", "json"],
                    "output_dir": out_dir,
                },
            )
            import time

            deadline = time.time() + 30.0
            job = None
            while time.time() < deadline:
                job = manager.get_job(job_id)
                if job and job.get("status") in ("done", "error"):
                    break
                time.sleep(0.2)

            self.assertIsNotNone(job)
            self.assertEqual(job["status"], "done")
            self.assertEqual(job["summary"]["ok"], 1)
            self.assertEqual(job["summary"]["failed"], 0)
            item = job["items"][0]
            self.assertEqual(item["status"], "ok")
            self.assertIn("pdf", item["outputs"])
            self.assertIn("docx", item["outputs"])
            self.assertIn("html", item["outputs"])
            self.assertTrue(os.path.isfile(item["outputs"]["pdf"]))
        finally:
            shutil.rmtree(out_dir, ignore_errors=True)


def main() -> int:
    suite = unittest.defaultTestLoader.loadTestsFromTestCase(TestEndToEnd)
    result = unittest.TextTestRunner(verbosity=2).run(suite)
    return 0 if result.wasSuccessful() else 1


if __name__ == "__main__":
    sys.exit(main())
