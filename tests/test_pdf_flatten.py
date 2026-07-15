"""Regressão do achatamento de pseudo-tabelas do PDF→DOCX (correção B3).

O pdf2docx envolve parágrafos corridos de PDFs nativos em tabelas (uma
palavra por célula), quebrando a visualização do redline. _flatten_pseudo_tables
desfaz isso preservando as tabelas reais.
"""
from __future__ import annotations

import os
import shutil
import tempfile
import unittest

from docx import Document

from app.extract.pdf_to_docx import _flatten_pseudo_tables, _is_pseudo_table


class FlattenPseudoTableTests(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.mkdtemp(prefix="comparedocs-flatten-")

    def tearDown(self) -> None:
        shutil.rmtree(self.tmp, ignore_errors=True)

    def _save(self, doc) -> str:
        path = os.path.join(self.tmp, "d.docx")
        doc.save(path)
        return path

    def _all_text(self, doc) -> str:
        return " ".join(" ".join(p.text.split()) for p in doc.paragraphs)

    def test_single_row_pseudo_table_is_flattened(self) -> None:
        doc = Document()
        doc.add_paragraph("Parágrafo normal antes.")
        t = doc.add_table(rows=1, cols=5)
        for cell, w in zip(
            t.rows[0].cells, ["CONSIDERANDO", "que", "a", "CONTRATANTE", "necessita"]
        ):
            cell.text = w
        self.assertTrue(_is_pseudo_table(t))
        path = self._save(doc)

        self.assertEqual(_flatten_pseudo_tables(path), 1)
        d2 = Document(path)
        self.assertEqual(len(d2.tables), 0)
        self.assertIn("CONSIDERANDO que a CONTRATANTE necessita", self._all_text(d2))

    def test_real_multirow_table_is_preserved(self) -> None:
        doc = Document()
        t = doc.add_table(rows=3, cols=3)
        for cell, w in zip(t.rows[0].cells, ["Parcela", "Vencimento", "Valor"]):
            cell.text = w
        t.rows[1].cells[0].text = "1"
        t.rows[1].cells[1].text = "01/2025"
        t.rows[1].cells[2].text = "R$ 55.500,00"
        self.assertFalse(_is_pseudo_table(t))
        path = self._save(doc)

        self.assertEqual(_flatten_pseudo_tables(path), 0)
        self.assertEqual(len(Document(path).tables), 1)

    def test_wide_multirow_pseudo_table_flattened_and_deduped(self) -> None:
        """2×N com palavra por célula (parágrafo que quebrou em 2 linhas) e
        células duplicadas pelo pdf2docx — achata e remove os duplicados."""
        doc = Document()
        t = doc.add_table(rows=2, cols=6)
        row0 = ["Confidenciais", "a", "que", "tiverem", "tiverem", "acesso"]
        for cell, w in zip(t.rows[0].cells, row0):
            cell.text = w
        row1 = ["em", "razão", "deste", "deste", "Contrato", "hoje"]
        for cell, w in zip(t.rows[1].cells, row1):
            cell.text = w
        self.assertTrue(_is_pseudo_table(t))
        path = self._save(doc)

        self.assertEqual(_flatten_pseudo_tables(path), 1)
        d2 = Document(path)
        self.assertEqual(len(d2.tables), 0)
        text = self._all_text(d2)
        self.assertIn("Confidenciais a que tiverem acesso", text)
        self.assertIn("em razão deste Contrato", text)
        self.assertNotIn("tiverem tiverem", text)
        self.assertNotIn("deste deste", text)


if __name__ == "__main__":
    unittest.main()
