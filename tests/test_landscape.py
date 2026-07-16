"""Regressão: documentos paisagem preservam geometria na prévia e no PDF."""
from __future__ import annotations

import unittest
from unittest import mock

from app.engine.compare import _preview_layout_from_doc
from app.models import Document
from app.output.redline_pdf import _pagesize_for_result
from reportlab.lib.pagesizes import A4, landscape


class LandscapeLayoutTests(unittest.TestCase):
    def test_preview_layout_marks_landscape(self) -> None:
        compare = Document(
            source_path="a.pdf",
            fmt="pdf",
            page_width_pt=842.0,
            page_height_pt=595.0,
        )
        base = Document(source_path="b.pdf", fmt="pdf")
        layout = _preview_layout_from_doc(compare, base)
        self.assertEqual(layout["orientation"], "landscape")
        self.assertEqual(layout["page_width_pt"], 842.0)
        self.assertEqual(layout["page_height_pt"], 595.0)

    def test_preview_layout_portrait_default(self) -> None:
        doc = Document(
            source_path="a.docx",
            fmt="docx",
            page_width_pt=595.0,
            page_height_pt=842.0,
        )
        layout = _preview_layout_from_doc(doc, doc)
        self.assertEqual(layout["orientation"], "portrait")

    def test_pagesize_for_result_uses_landscape(self) -> None:
        from app.models import ComparisonResult, Stats

        result = ComparisonResult(
            base_path="a",
            compare_path="b",
            stats=Stats(),
            preview_layout={
                "orientation": "landscape",
                "page_width_pt": 842.0,
                "page_height_pt": 595.0,
            },
        )
        size = _pagesize_for_result(result)
        self.assertGreater(size[0], size[1])

    def test_pagesize_default_a4(self) -> None:
        from app.models import ComparisonResult, Stats

        result = ComparisonResult(
            base_path="a", compare_path="b", stats=Stats(), preview_layout=None
        )
        self.assertEqual(_pagesize_for_result(result), A4)


class ApplyPdfGeometryTests(unittest.TestCase):
    def test_apply_sets_landscape_on_docx(self) -> None:
        import os
        import tempfile

        from docx import Document as DocxDocument
        from docx.enum.section import WD_ORIENT
        from docx.shared import Pt

        from app.extract.pdf_to_docx import _apply_pdf_page_geometry

        tmp = tempfile.mkdtemp(prefix="comparedocs-land-")
        try:
            docx_path = os.path.join(tmp, "out.docx")
            d = DocxDocument()
            d.add_paragraph("hello")
            # Começa em A4 retrato
            sec = d.sections[0]
            sec.orientation = WD_ORIENT.PORTRAIT
            sec.page_width = Pt(595)
            sec.page_height = Pt(842)
            d.save(docx_path)

            with mock.patch(
                "app.extract.pdf_to_docx._pdf_page_size_pt",
                return_value=(842.0, 595.0),
            ):
                _apply_pdf_page_geometry("fake.pdf", docx_path)

            out = DocxDocument(docx_path)
            sec = out.sections[0]
            self.assertGreater(float(sec.page_width.pt), float(sec.page_height.pt))
        finally:
            import shutil

            shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
